import io
import json
import os
import tempfile
import datetime as dt
from functools import lru_cache
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

import pandas as pd
import streamlit as st
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

try:
    pd.set_option("future.no_silent_downcasting", True)
except Exception:
    pass

try:
    from google.cloud import storage
except ImportError:  # pragma: no cover - optional dependency for remote archives
    storage = None

from generate_kristin_robbins_votes import WORKBOOK_HEADERS, write_workbook
from json_legiscan_loader import (
    collect_legislator_names_json,
    determine_json_state,
    extract_archives,
    gather_json_session_dirs,
)
from json_vote_builder import STATUS_LABELS, collect_vote_rows_from_json, extract_crossfile_fields


def _settings_value(key: str, default: Optional[str] = None) -> Optional[str]:
    secrets_obj = getattr(st, "secrets", None)
    if secrets_obj and key in secrets_obj:
        return secrets_obj[key]
    return os.environ.get(key, default)


def _service_account_info() -> Optional[Dict[str, object]]:
    secrets_obj = getattr(st, "secrets", None)
    if not secrets_obj or "gcp_service_account" not in secrets_obj:
        return None
    raw_info = secrets_obj["gcp_service_account"]
    if hasattr(raw_info, "to_dict"):
        return raw_info.to_dict()
    if isinstance(raw_info, dict):
        return raw_info.copy()
    try:
        return dict(raw_info)
    except TypeError:
        return None


DEFAULT_JSON_DATA_DIR = Path(__file__).resolve().parent / "JSON DATA"
GCS_BUCKET_NAME = _settings_value("ILLUMIS_GCS_BUCKET")
GCS_MANIFEST_BLOB = _settings_value("ILLUMIS_GCS_MANIFEST", "manifest.json")
archive_cache_setting = _settings_value(
    "ILLUMIS_ARCHIVE_CACHE_DIR",
    str(Path(tempfile.gettempdir()) / "legiscan-json-archives"),
)
ARCHIVE_CACHE_DIR = Path(archive_cache_setting)
ARCHIVE_CACHE_DIR.mkdir(parents=True, exist_ok=True)

service_account = _service_account_info()
if service_account:
    creds_path = ARCHIVE_CACHE_DIR / "gcs-creds.json"
    creds_path.write_text(json.dumps(service_account))
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = str(creds_path)
    project_id = service_account.get("project_id")
    if project_id:
        os.environ.setdefault("GOOGLE_CLOUD_PROJECT", project_id)

USE_REMOTE_ARCHIVES = bool(GCS_BUCKET_NAME)
JSON_DATA_DIR = DEFAULT_JSON_DATA_DIR if not USE_REMOTE_ARCHIVES else ARCHIVE_CACHE_DIR

SESSION_CACHE_KEY = "json_vote_summary"
ALL_STATES_LABEL = "All States"


@lru_cache(maxsize=1)
def _storage_client():
    if storage is None:
        raise RuntimeError(
            "google-cloud-storage is required when ILLUMIS_GCS_BUCKET is configured."
        )
    return storage.Client()


def _parse_iso8601(value: Optional[str]) -> Optional[dt.datetime]:
    if not value:
        return None
    text = value.strip()
    if not text:
        return None
    text = text.replace("Z", "+00:00")
    try:
        return dt.datetime.fromisoformat(text)
    except ValueError:
        return None


def _normalize_manifest_entry(state_code: str, entry: object) -> Optional[Dict[str, object]]:
    if isinstance(entry, str):
        blob_path = entry
        name = Path(entry).name
        updated = None
        size = None
    elif isinstance(entry, dict):
        blob_path = (
            entry.get("blob_path")
            or entry.get("blob")
            or entry.get("path")
            or entry.get("name")
        )
        name = entry.get("name") or (Path(blob_path).name if blob_path else None)
        updated = entry.get("updated") or entry.get("last_modified")
        size = entry.get("size")
    else:
        return None
    if not blob_path or not name:
        return None
    return {
        "state": state_code,
        "name": name,
        "blob_path": blob_path,
        "updated": updated,
        "size": size,
    }


def _build_manifest(payload: object) -> Dict[str, List[Dict[str, object]]]:
    manifest: Dict[str, List[Dict[str, object]]] = {}
    if isinstance(payload, dict):
        items = payload.items()
    elif isinstance(payload, list):
        items = (
            ((entry.get("state") or "").strip().upper(), [entry])
            for entry in payload
            if isinstance(entry, dict)
        )
    else:
        raise ValueError(
            "GCS manifest must be a JSON object keyed by state codes or a list of state entries."
        )
    for state_key, entries in items:
        code = (state_key or "").strip().upper()
        if not code:
            continue
        normalized_entries: List[Dict[str, object]] = []
        if not isinstance(entries, list):
            entries = [entries]
        for entry in entries:
            normalized = _normalize_manifest_entry(code, entry)
            if normalized:
                normalized_entries.append(normalized)
        if normalized_entries:
            manifest[code] = normalized_entries
    return manifest


@lru_cache(maxsize=1)
def _cached_manifest() -> Dict[str, List[Dict[str, object]]]:
    if not USE_REMOTE_ARCHIVES:
        return {}
    client = _storage_client()
    bucket = client.bucket(GCS_BUCKET_NAME)
    blob = bucket.blob(GCS_MANIFEST_BLOB)
    payload = json.loads(blob.download_as_text())
    return _build_manifest(payload)


@lru_cache(maxsize=1)
def _remote_archive_lookup() -> Dict[str, Dict[str, object]]:
    lookup: Dict[str, Dict[str, object]] = {}
    for entries in _cached_manifest().values():
        for entry in entries:
            lookup[entry["name"]] = entry
    return lookup


def _ensure_remote_archive(entry: Dict[str, object]) -> Path:
    ARCHIVE_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    local_path = ARCHIVE_CACHE_DIR / str(entry["name"])
    updated = _parse_iso8601(entry.get("updated"))
    if local_path.exists() and updated:
        normalized_updated = (
            updated if updated.tzinfo else updated.replace(tzinfo=dt.timezone.utc)
        )
        updated_ts = normalized_updated.timestamp()
        if local_path.stat().st_mtime >= updated_ts:
            return local_path
    elif local_path.exists():
        return local_path
    client = _storage_client()
    bucket = client.bucket(GCS_BUCKET_NAME)
    blob = bucket.blob(str(entry["blob_path"]))
    tmp_path = local_path.with_suffix(local_path.suffix + ".tmp")
    blob.download_to_filename(tmp_path)
    tmp_path.replace(local_path)
    if updated:
        normalized_updated = (
            updated if updated.tzinfo else updated.replace(tzinfo=dt.timezone.utc)
        )
        timestamp = normalized_updated.timestamp()
        os.utime(local_path, (timestamp, timestamp))
    return local_path
STATE_CHOICES = [
    ("Alabama", "AL"),
    ("Alaska", "AK"),
    ("Arizona", "AZ"),
    ("Arkansas", "AR"),
    ("California", "CA"),
    ("Colorado", "CO"),
    ("Connecticut", "CT"),
    ("Delaware", "DE"),
    ("Florida", "FL"),
    ("Georgia", "GA"),
    ("Hawaii", "HI"),
    ("Idaho", "ID"),
    ("Illinois", "IL"),
    ("Indiana", "IN"),
    ("Iowa", "IA"),
    ("Kansas", "KS"),
    ("Kentucky", "KY"),
    ("Louisiana", "LA"),
    ("Maine", "ME"),
    ("Maryland", "MD"),
    ("Massachusetts", "MA"),
    ("Michigan", "MI"),
    ("Minnesota", "MN"),
    ("Mississippi", "MS"),
    ("Missouri", "MO"),
    ("Montana", "MT"),
    ("Nebraska", "NE"),
    ("Nevada", "NV"),
    ("New Hampshire", "NH"),
    ("New Jersey", "NJ"),
    ("New Mexico", "NM"),
    ("New York", "NY"),
    ("North Carolina", "NC"),
    ("North Dakota", "ND"),
    ("Ohio", "OH"),
    ("Oklahoma", "OK"),
    ("Oregon", "OR"),
    ("Pennsylvania", "PA"),
    ("Rhode Island", "RI"),
    ("South Carolina", "SC"),
    ("South Dakota", "SD"),
    ("Tennessee", "TN"),
    ("Texas", "TX"),
    ("Utah", "UT"),
    ("Vermont", "VT"),
    ("Virginia", "VA"),
    ("Washington", "WA"),
    ("West Virginia", "WV"),
    ("Wisconsin", "WI"),
    ("Wyoming", "WY"),
]
STATE_NAME_TO_CODE = {name: code for name, code in STATE_CHOICES}
STATE_CODE_TO_NAME = {code: name for name, code in STATE_CHOICES}
FOCUS_PARTY_LOOKUP = {
    "Legislator's Party": None,
    "Democrat": "Democrat",
    "Republican": "Republican",
    "Independent": "Other",
}
FILTER_OPTIONS = [
    "All Votes",
    "Votes Against Party",
    "Minority Votes",
    "Deciding Votes",
    "Skipped Votes",
    "Search By Term",
    "Sponsored/Cosponsored Bills",
]
WORKBOOK_VIEWS = [
    "All Votes",
    "Votes Against Party",
    "Minority Votes",
    "Deciding Votes",
    "Skipped Votes",
    "Sponsored/Cosponsored Bills",
]


def _resolve_archives(selected_names: List[str]) -> List[Path]:
    if USE_REMOTE_ARCHIVES:
        lookup = _remote_archive_lookup()
        missing = [name for name in selected_names if name not in lookup]
        if missing:
            raise FileNotFoundError(f"Missing archive(s) in manifest: {', '.join(missing)}")
        return [_ensure_remote_archive(lookup[name]) for name in selected_names]

    lookup = {path.name: path for path in DEFAULT_JSON_DATA_DIR.glob("*.zip")}
    missing = [name for name in selected_names if name not in lookup]
    if missing:
        raise FileNotFoundError(f"Missing archive(s): {', '.join(missing)}")
    return [lookup[name] for name in selected_names]


@st.cache_data(show_spinner=False)
def load_legislators_for_archives(archive_names: Tuple[str, ...]) -> Tuple[List[str], Optional[str]]:
    selected_paths = [Path(name) for name in archive_names]
    extracted = extract_archives(selected_paths)
    try:
        base_dirs = [item.base_path for item in extracted]
        session_dirs = gather_json_session_dirs(base_dirs)
        dataset_state = determine_json_state(session_dirs)
        names = collect_legislator_names_json(session_dirs)
        return names, dataset_state
    finally:
        for item in extracted:
            item.cleanup()


def _load_legislator_names(selected_paths: List[Path]) -> Tuple[List[str], Optional[str]]:
    extracted = extract_archives(selected_paths)
    try:
        base_dirs = [item.base_path for item in extracted]
        session_dirs = gather_json_session_dirs(base_dirs)
        dataset_state = determine_json_state(session_dirs)
        names = collect_legislator_names_json(session_dirs)
        return names, dataset_state
    finally:
        for item in extracted:
            item.cleanup()


def _build_vote_rows(selected_paths: List[Path], legislator_name: str) -> List[List]:
    extracted = extract_archives(selected_paths)
    try:
        base_dirs = [item.base_path for item in extracted]
        session_dirs = gather_json_session_dirs(base_dirs)
        return collect_vote_rows_from_json(session_dirs, legislator_name)
    finally:
        for item in extracted:
            item.cleanup()


def _prepare_dataframe(rows: List[List]) -> pd.DataFrame:
    df = pd.DataFrame(rows, columns=WORKBOOK_HEADERS)
    df["Date_dt"] = pd.to_datetime(df["Date"], errors="coerce")
    df["Year"] = df["Date_dt"].dt.year.astype("Int64")
    return df


def _render_state_filter() -> Tuple[List[str], List[str]]:
    st.sidebar.header("Dataset Selection")
    selection = st.sidebar.selectbox(
        "State",
        options=[name for name, _ in STATE_CHOICES],
        index=None,
        placeholder="Select a state",
        help="Choose a single state to load its JSON archives.",
    )
    if selection:
        return [selection], [STATE_NAME_TO_CODE[selection]]
    return [], []


def _collect_archives_for_states(state_codes: List[str]) -> List[str]:
    if not state_codes:
        return []

    if USE_REMOTE_ARCHIVES:
        manifest = _cached_manifest()
        if not manifest:
            st.error(
                "No JSON archives are listed in the configured GCS manifest. "
                "Confirm ILLUMIS_GCS_BUCKET and ILLUMIS_GCS_MANIFEST are correct."
            )
            st.stop()
        selected: List[str] = []
        for code in state_codes:
            selected.extend([entry["name"] for entry in manifest.get(code, [])])
        return selected

    available_archives = sorted(DEFAULT_JSON_DATA_DIR.glob("*.zip"))
    if not available_archives:
        st.error(f"No JSON ZIP archives found in {DEFAULT_JSON_DATA_DIR}.")
        st.stop()

    selected: List[str] = []
    for path in available_archives:
        code = path.name[:2].upper()
        if code in state_codes:
            selected.append(path.name)
    return selected


PARTY_CODE_MAP = {
    "D": "Democrat",
    "DEM": "Democrat",
    "DFL": "Democrat",
    "R": "Republican",
    "REP": "Republican",
    "GOP": "Republican",
}


def _normalize_party_label(party_code: Optional[str]) -> str:
    code = (party_code or "").strip().upper()
    if not code:
        return ""
    if code in PARTY_CODE_MAP:
        return PARTY_CODE_MAP[code]
    if code in {"I", "IND", "IND.", "INDP", "INDEPENDENT"}:
        return "Other"
    return "Other"


def _format_json_date(date_str: Optional[str]) -> str:
    if not date_str:
        return ""
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%Y/%m/%d"):
        try:
            return dt.datetime.strptime(date_str, fmt).strftime("%m/%d/%Y")
        except ValueError:
            continue
    return date_str


def _format_roll_details_json(vote_entry: Optional[dict]) -> str:
    if not vote_entry:
        return ""
    desc = (vote_entry.get("desc") or "").strip()
    yea = safe_int(vote_entry.get("yea"))
    nay = safe_int(vote_entry.get("nay"))
    suffix = ""
    if yea or nay:
        suffix = f" ({yea}-Y {nay}-N)"
    return f"{desc}{suffix}" if desc else suffix.strip()


def _normalize_chamber_label(value: Optional[str]) -> str:
    token = (value or "").strip()
    if not token:
        return "Chamber"
    upper = token.upper()
    if upper in {"H", "HOUSE", "LOWER"}:
        return "House"
    if upper in {"S", "SENATE", "UPPER"}:
        return "Senate"
    return token


def _parse_bill_id_list(raw_text: str) -> List[str]:
    if not raw_text:
        return []
    normalized: List[str] = []
    for part in raw_text.replace("\r", "\n").replace(",", "\n").split("\n"):
        value = part.strip()
        if value:
            normalized.append(value)
    return normalized


def _latest_history_entry_json(history: Optional[List[dict]]) -> Tuple[str, str]:
    if not history:
        return "", ""
    latest = history[-1]
    action = (latest.get("action") or "").strip()
    date_str = _format_json_date(latest.get("date"))
    return action, date_str


def _build_json_sponsor_metadata(bill: dict, status: str) -> Dict[str, object]:
    session_meta = bill.get("session") or {}
    session_label = session_meta.get("session_name") or bill.get("session_id") or ""
    bill_number = str(bill.get("bill_number") or "")
    bill_id = bill.get("bill_id") or ""
    title = (bill.get("title") or "").strip()
    description = (bill.get("description") or "").strip()
    bill_motion = description or title or bill_number
    url = bill.get("state_link") or bill.get("url") or ""
    status_code = bill.get("status") or ""
    status_desc = STATUS_LABELS.get(int(status_code), str(status_code)) if status_code else ""
    status_date = _format_json_date(bill.get("status_date"))
    last_action, last_action_date = _latest_history_entry_json(bill.get("history") or [])
    votes = bill.get("votes") or []
    first_vote = votes[0] if votes else None
    roll_call_id = first_vote.get("roll_call_id") if first_vote else ""
    roll_details = _format_roll_details_json(first_vote)
    roll_date = _format_json_date(first_vote.get("date") if first_vote else "")
    result = 1 if first_vote and first_vote.get("passed") else 0
    chamber_value = ""
    if first_vote:
        chamber_token = (first_vote.get("chamber") or "").upper()
        if chamber_token.startswith("H"):
            chamber_value = "House"
        elif chamber_token.startswith("S"):
            chamber_value = "Senate"
    if not chamber_value:
        body_token = (bill.get("body") or "").upper()
        if body_token == "H":
            chamber_value = "House"
        elif body_token == "S":
            chamber_value = "Senate"
    excel_date = roll_date or status_date or last_action_date
    crossfile_id, crossfile_number = extract_crossfile_fields(bill)
    return {
        "session_id": session_label,
        "bill_number": bill_number,
        "bill_id": bill_id,
        "crossfile_bill_id": crossfile_id,
        "crossfile_bill_number": crossfile_number,
        "bill_title": title,
        "bill_description": description,
        "bill_motion": bill_motion,
        "bill_url": url,
        "status_code": status_code,
        "status_desc": status_desc,
        "status_date": status_date,
        "last_action": last_action,
        "last_action_date": last_action_date,
        "roll_call_id": roll_call_id,
        "roll_details": roll_details,
        "roll_date": roll_date,
        "excel_date": excel_date,
        "result": result,
        "chamber": chamber_value,
        "sponsorship_status": status,
    }


def _sanitize_sheet_title(title: str, used_titles: set[str]) -> str:
    cleaned = "".join(ch if ch not in '[]:*?/\\' else "_" for ch in title).strip() or "Sheet"
    cleaned = cleaned[:31]
    base = cleaned
    counter = 1
    while cleaned in used_titles:
        suffix = f"_{counter}"
        if len(base) + len(suffix) > 31:
            cleaned = base[: 31 - len(suffix)] + suffix
        else:
            cleaned = base + suffix
        counter += 1
    used_titles.add(cleaned)
    return cleaned


def write_multi_sheet_workbook(
    sheet_specs: List[Tuple[str, List[str], List[List]]],
    output: io.BytesIO,
) -> None:
    wb = Workbook()
    used_titles: set[str] = set()
    first_sheet = True
    for sheet_name, headers, rows in sheet_specs:
        title = _sanitize_sheet_title(sheet_name, used_titles)
        if first_sheet:
            ws = wb.active
            ws.title = title
            first_sheet = False
        else:
            ws = wb.create_sheet(title=title)
        ws.append(headers)
        for row in rows:
            ws.append(row)
    wb.save(output)


def _collect_sponsor_lookup_json(
    archive_paths: List[Path],
    legislator_name: str,
) -> Tuple[
    Dict[Tuple[str, str], str],
    Dict[Tuple[str, str], Dict[str, object]],
    str,
    Dict[Tuple[str, str], List[Dict[str, object]]],
]:
    if not archive_paths:
        return {}, {}, "", {}
    extracted = extract_archives(archive_paths)
    try:
        base_dirs = [item.base_path for item in extracted]
        session_dirs = gather_json_session_dirs(base_dirs)
        sponsor_lookup: Dict[Tuple[str, str], str] = {}
        sponsor_metadata: Dict[Tuple[str, str], Dict[str, object]] = {}
        bill_vote_metadata: Dict[Tuple[str, str], List[Dict[str, object]]] = {}
        legislator_party_label = ""
        target = legislator_name.strip()
        for session_dir in session_dirs:
            bill_dir = Path(session_dir) / "bill"
            if not bill_dir.exists():
                continue
            for bill_path in bill_dir.glob("*.json"):
                with bill_path.open(encoding="utf-8") as fh:
                    data = json.load(fh)
                bill = data.get("bill") or {}
                session_meta = bill.get("session") or {}
                session_label = session_meta.get("session_name") or bill.get("session_id") or ""
                bill_number = str(bill.get("bill_number") or "").strip()
                if not session_label or not bill_number:
                    continue
                bill_id = str(bill.get("bill_id") or "").strip()
                if not bill_id:
                    continue
                key = (session_label, bill_id)
                bill_vote_metadata[key] = bill.get("votes") or []
                for sponsor in bill.get("sponsors") or []:
                    name = (sponsor.get("name") or "").strip()
                    if name != target:
                        continue
                    sponsor_type_id = sponsor.get("sponsor_type_id")
                    status = "Primary Sponsor" if sponsor_type_id == 1 else "Cosponsor"
                    existing = sponsor_lookup.get(key)
                    if existing == "Primary Sponsor" and status != "Primary Sponsor":
                        continue
                    if status == "Primary Sponsor" or key not in sponsor_lookup:
                        sponsor_lookup[key] = status
                    if not legislator_party_label:
                        legislator_party_label = _normalize_party_label(sponsor.get("party"))
                    meta = sponsor_metadata.get(key)
                    new_meta = _build_json_sponsor_metadata(bill, status)
                    if meta is None or status == "Primary Sponsor":
                        sponsor_metadata[key] = new_meta
                    elif meta.get("sponsorship_status") != "Primary Sponsor":
                        sponsor_metadata[key]["sponsorship_status"] = status
        return sponsor_lookup, sponsor_metadata, legislator_party_label, bill_vote_metadata
    finally:
        for item in extracted:
            item.cleanup()


def _create_sponsor_only_rows(
    sponsor_metadata: Dict[Tuple[str, str], Dict[str, object]],
    existing_keys: Set[Tuple[str, str]],
    legislator_name: str,
    legislator_party_label: str,
) -> List[Dict[str, object]]:
    rows: List[Dict[str, object]] = []
    count_start_idx = WORKBOOK_HEADERS.index("Democrat_For")
    party_label = legislator_party_label or "Other"
    for idx, (key, meta) in enumerate(sponsor_metadata.items()):
        if key in existing_keys:
            continue
        session_label, sponsor_bill_id = key
        roll_call_value = meta.get("roll_call_id", "")
        try:
            normalized_roll_id = int(str(roll_call_value))
        except (TypeError, ValueError):
            normalized_roll_id = -(10**9 + idx)
        row = {header: "" for header in WORKBOOK_HEADERS}
        for header in WORKBOOK_HEADERS[count_start_idx:]:
            row[header] = 0
        row.update(
            {
                "Chamber": meta.get("chamber", ""),
                "Session": meta.get("session_id", ""),
                "Bill Number": meta.get("bill_number", ""),
                "Bill ID": meta.get("bill_id") or sponsor_bill_id or "",
                "Cross-file Bill ID": meta.get("crossfile_bill_id", ""),
                "Cross-file Bill Number": meta.get("crossfile_bill_number", ""),
                "Bill Motion": meta.get("bill_motion", "") or meta.get("bill_title", ""),
                "URL": meta.get("bill_url", ""),
                "Bill Title": meta.get("bill_title", ""),
                "Bill Description": meta.get("bill_description", "") or meta.get("bill_title", ""),
                "Roll Details": meta.get("roll_details", ""),
                "Roll Call ID": normalized_roll_id,
                "Last Action Date": meta.get("last_action_date", ""),
                "Last Action": meta.get("last_action", ""),
                "Status": meta.get("status_code", ""),
                "Status Description": meta.get("status_desc", ""),
                "Status Date": meta.get("status_date", ""),
                "Person": legislator_name,
                "Person Party": party_label,
                "Date": meta.get("excel_date", ""),
                "Result": meta.get("result", ""),
            }
        )
        rows.append(
            {
                **row,
                "Sponsorship Status": meta.get("sponsorship_status", ""),
            }
        )
    return rows


def _build_legislator_dataset(
    selected_paths: List[Path],
    legislator_name: str,
) -> Dict[str, object]:
    rows = _build_vote_rows(selected_paths, legislator_name)
    summary_df = _prepare_dataframe(rows)
    (
        sponsor_lookup,
        sponsor_metadata,
        legislator_party_label,
        bill_vote_metadata,
    ) = _collect_sponsor_lookup_json(
        selected_paths,
        legislator_name,
    )
    session_series = summary_df["Session"].astype(str)
    bill_id_series = summary_df["Bill ID"].astype(str)
    summary_df["Sponsorship Status"] = [
        sponsor_lookup.get((session, bill_id), "")
        for session, bill_id in zip(session_series, bill_id_series)
    ]
    return {
        "rows": rows,
        "df": summary_df,
        "sponsor_metadata": sponsor_metadata,
        "legislator_party_label": legislator_party_label,
        "bill_vote_metadata": bill_vote_metadata,
    }


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    if not url:
        paragraph.add_run(text)
        return
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _format_docx_date(date_value: str) -> str:
    text = (date_value or "").strip()
    if not text:
        return text
    date_formats = ("%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y")
    for fmt in date_formats:
        try:
            parsed = dt.datetime.strptime(text, fmt)
            return f"{parsed.month}/{parsed.day}/{str(parsed.year)[-2:]}"
        except ValueError:
            continue
    try:
        parsed = dt.datetime.fromisoformat(text)
        return f"{parsed.month}/{parsed.day}/{str(parsed.year)[-2:]}"
    except ValueError:
        return text


def _resolve_vote_phrases(vote_bucket: str) -> Tuple[str, str]:
    bucket = (vote_bucket or "").strip().lower()
    if bucket == "for":
        return "VOTED FOR", "voted for"
    if bucket == "against":
        return "VOTED AGAINST", "voted against"
    if bucket == "absent":
        return "WAS ABSENT FOR", "was absent for"
    return "DID NOT VOTE ON", "did not vote on"


def _build_json_bullet_summary_doc(
    rows: pd.DataFrame,
    legislator_name: str,
    filter_label: str,
    state_label: str,
    *,
    bill_vote_metadata: Optional[Dict[Tuple[str, str], List[Dict[str, object]]]] = None,
    full_summary_df: Optional[pd.DataFrame] = None,
    include_amendments: bool = False,
) -> io.BytesIO:
    bill_vote_metadata = bill_vote_metadata or {}
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)

    doc.add_heading(f"{legislator_name} - {filter_label}", level=1)

    if rows.empty:
        doc.add_paragraph("No records available for this selection.")
    else:
        working_rows = rows.copy()
        if "Date_dt" not in working_rows.columns or working_rows["Date_dt"].isna().any():
            working_rows["Date_dt"] = pd.to_datetime(
                working_rows.get("Date"), errors="coerce"
            )
        working_rows = working_rows.sort_values(
            by=["Date_dt", "Roll Call ID"] if "Roll Call ID" in working_rows.columns else ["Date_dt"],
            kind="mergesort",
            na_position="last",
        )

        bill_rows_map: Dict[Tuple[str, str], List[pd.Series]] = {}
        bill_context_row: Dict[Tuple[str, str], pd.Series] = {}
        ordered_bill_keys: List[Tuple[str, str]] = []
        for _, row in working_rows.iterrows():
            session_label = str(row.get("Session") or "").strip()
            bill_id = str(row.get("Bill ID") or "").strip()
            if not session_label or not bill_id:
                continue
            key = (session_label, bill_id)
            bill_rows_map.setdefault(key, []).append(row)
            if key not in bill_context_row:
                bill_context_row[key] = row
                ordered_bill_keys.append(key)

        vote_row_lookup: Dict[int, pd.Series] = {}
        source_df = full_summary_df if full_summary_df is not None else working_rows
        if source_df is not None:
            for _, src_row in source_df.iterrows():
                roll_call_identifier = safe_int(src_row.get("Roll Call ID"))
                if roll_call_identifier:
                    vote_row_lookup[roll_call_identifier] = src_row

        for key in ordered_bill_keys:
            context_row = bill_context_row.get(key)
            if context_row is None:
                continue
            bill_number = (context_row.get("Bill Number") or "").strip() or "Unknown bill"
            bill_title = (context_row.get("Bill Title") or "").strip()
            bill_description = (context_row.get("Bill Description") or bill_title or "").strip()
            if bill_description:
                bill_description = bill_description.rstrip(".")
            sponsorship = (context_row.get("Sponsorship Status") or "").strip()
            last_action = (context_row.get("Last Action") or context_row.get("Status Description") or "").strip()
            roll_call_entries = bill_vote_metadata.get(key) or []
            if not include_amendments:
                roll_call_entries = [
                    entry
                    for entry in roll_call_entries
                    if not _is_amendment_roll_call_entry(entry)
                ]

            bill_rows = bill_rows_map.get(key, [])
            non_amendment_rows = bill_rows
            if not include_amendments:
                non_amendment_rows = [
                    row for row in bill_rows if not _is_amendment_row(row)
                ]

            if not roll_call_entries and not non_amendment_rows:
                continue

            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(12)
            paragraph.paragraph_format.space_before = Pt(0)

            header_row = None
            if roll_call_entries:
                for entry in roll_call_entries:
                    rcid = safe_int(entry.get("roll_call_id"))
                    vote_row = vote_row_lookup.get(rcid)
                    if vote_row is not None:
                        header_row = vote_row
                        break
            if header_row is None:
                header_row_source = non_amendment_rows or bill_rows or [context_row]
                header_row = header_row_source[0]
            example_row = header_row
            vote_dt = example_row.get("Date_dt")
            if pd.isna(vote_dt):
                display_date = (example_row.get("Date") or "").strip()
            else:
                display_date = vote_dt.strftime("%B %d, %Y")
            vote_bucket = example_row.get("Vote Bucket", "")
            vote_upper, vote_lower = _resolve_vote_phrases(vote_bucket)
            first_sentence = (
                f"{display_date or 'Date unknown'}: {legislator_name} "
                f"{vote_upper.title()} {bill_number}"
            )
            if bill_title:
                first_sentence += f" - {bill_title}"
            if sponsorship:
                first_sentence += f" ({sponsorship})"
            first_sentence += "."
            paragraph.add_run(first_sentence + " ").bold = True

            if bill_description:
                paragraph.add_run(f"{legislator_name} {vote_lower} {bill_number}: \"{bill_description}.\" ")

            summary_sentences: List[str] = []

            if not roll_call_entries:
                for row in non_amendment_rows:
                    summary_sentences.append(
                        _build_roll_call_sentence_from_row(
                            row,
                            None,
                            legislator_name,
                            bill_number,
                        )
                    )
            else:
                sorted_roll_calls = sorted(
                    roll_call_entries,
                    key=lambda entry: _parse_roll_call_datetime(entry.get("date")),
                )
                for entry in sorted_roll_calls:
                    roll_call_id = safe_int(entry.get("roll_call_id"))
                    vote_row = vote_row_lookup.get(roll_call_id)
                    summary_sentences.append(
                        _build_roll_call_sentence_from_entry(
                            entry,
                            vote_row,
                            legislator_name,
                            bill_number,
                        )
                    )

            summary_sentences = [s for s in summary_sentences if s]
            if summary_sentences:
                paragraph.add_run(" ".join(summary_sentences) + " ")

            if last_action:
                paragraph.add_run(f"Latest action: {last_action}. ")

            paragraph.add_run("[")
            bracket_chamber = _normalize_chamber_label(context_row.get("Chamber"))
            paragraph.add_run(f"{state_label or 'State'} {bracket_chamber}, {bill_number}, ")
            date_bracket = display_date or "Date unknown"
            url = (context_row.get("URL") or "").strip()
            if url:
                _add_hyperlink(paragraph, url, date_bracket)
            else:
                paragraph.add_run(date_bracket)
            paragraph.add_run("]")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _parse_roll_call_datetime(value: Optional[str]) -> dt.datetime:
    if not value:
        return dt.datetime.min
    try:
        return dt.datetime.fromisoformat(value)
    except ValueError:
        pass
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y"):
        try:
            return dt.datetime.strptime(value, fmt)
        except ValueError:
            continue
    return dt.datetime.min


def _is_amendment_roll_call_entry(entry: Optional[Dict[str, object]]) -> bool:
    if not entry:
        return False
    desc = (entry.get("desc") or "").strip().lower()
    return "amendment" in desc


def _is_amendment_row(row: pd.Series) -> bool:
    details = str(row.get("Roll Details") or "").lower()
    return "amendment" in details


def _build_roll_call_sentence_from_row(
    row: pd.Series,
    entry: Optional[Dict[str, object]],
    legislator_name: str,
    bill_number: str,
) -> str:
    vote_dt = row.get("Date_dt")
    if pd.isna(vote_dt):
        display_date = (row.get("Date") or "").strip()
    else:
        display_date = vote_dt.strftime("%m/%d/%Y")
    chamber = _normalize_chamber_label(row.get("Chamber"))
    result_value = row.get("Result")
    vote_summary_text = _format_vote_summary_counts(row)
    if not vote_summary_text:
        vote_summary_text = _format_roll_call_counts_entry(entry)
    sentence = _format_roll_call_sentence(
        display_date,
        chamber,
        result_value,
        vote_summary_text,
        bill_number,
        entry,
    )
    vote_bucket = row.get("Vote Bucket", "")
    _, vote_lower = _resolve_vote_phrases(vote_bucket)
    if vote_lower:
        sentence += f", {legislator_name} {vote_lower}."
    else:
        sentence += "."
    return sentence


def _build_roll_call_sentence_from_entry(
    entry: Dict[str, object],
    vote_row: Optional[pd.Series],
    legislator_name: str,
    bill_number: str,
) -> str:
    display_date = _format_json_date(entry.get("date"))
    chamber = _normalize_chamber_label(
        entry.get("chamber") or (vote_row.get("Chamber") if vote_row is not None else "")
    )
    result_value = 1 if safe_int(entry.get("passed")) == 1 else 0
    vote_summary_text = _format_roll_call_counts_entry(entry)
    sentence = _format_roll_call_sentence(
        display_date,
        chamber,
        result_value,
        vote_summary_text,
        bill_number,
        entry,
    )
    if vote_row is not None:
        vote_bucket = vote_row.get("Vote Bucket", "")
        _, vote_lower = _resolve_vote_phrases(vote_bucket)
        sentence += f", {legislator_name} {vote_lower}."
    else:
        sentence += f" {legislator_name} did not cast a vote."
    return sentence

def safe_int(value: object) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return 0


def _format_vote_summary_counts(row: pd.Series) -> Optional[str]:
    required_cols = {"Total_For", "Total_Against", "Total_Not", "Total_Absent"}
    if not required_cols.issubset(row.index):
        return None
    total_for = safe_int(row.get("Total_For"))
    total_against = safe_int(row.get("Total_Against"))
    total_not = safe_int(row.get("Total_Not"))
    total_absent = safe_int(row.get("Total_Absent"))
    total_sum = total_for + total_against + total_not + total_absent
    if total_sum == 0:
        return None
    return f"{total_for}-{total_against}-{total_not}-{total_absent}"


def _format_roll_call_counts_entry(entry: Optional[Dict[str, object]]) -> Optional[str]:
    if not entry:
        return None
    total_for = safe_int(entry.get("yea"))
    total_against = safe_int(entry.get("nay"))
    total_not = safe_int(entry.get("nv"))
    total_absent = safe_int(entry.get("absent"))
    total_sum = total_for + total_against + total_not + total_absent
    if total_sum == 0:
        return None
    return f"{total_for}-{total_against}-{total_not}-{total_absent}"


def _format_roll_call_sentence(
    display_date: Optional[str],
    chamber: str,
    result_value: object,
    vote_summary_text: Optional[str],
    bill_number: str,
    roll_call_entry: Optional[Dict[str, object]],
) -> str:
    vote_date = display_date or "Date unknown"
    chamber_label = chamber or "Chamber"
    action = "pass" if safe_int(result_value) == 1 else "reject"
    sentence = f"On {vote_date} the {chamber_label} voted to {action} {bill_number}"
    if roll_call_entry:
        desc = (roll_call_entry.get("desc") or "").strip()
        if desc:
            if "amendment" in desc.lower():
                cleaned = desc.replace("Floor", "").strip()
                sentence += f", {cleaned}"
            else:
                sentence += f", {desc}"
    if vote_summary_text:
        sentence += f", ({vote_summary_text})"
    return sentence


def _format_outcome_sentence(
    display_date: str,
    chamber: str,
    bill_number: str,
    result_value: object,
    vote_summary_text: Optional[str],
) -> str:
    vote_date = display_date or "Date unknown"
    chamber_label = _normalize_chamber_label(chamber)
    normalized_bill = bill_number or "the bill"
    action = "pass" if safe_int(result_value) == 1 else "reject"
    sentence = f"On {vote_date} the {chamber_label} voted to {action} {normalized_bill}"
    if vote_summary_text:
        sentence += f", ({vote_summary_text})"
    sentence += "."
    return sentence


def _apply_deciding_vote_filter(df: pd.DataFrame, max_vote_diff: int) -> pd.Series:
    total_for = df["Total_For"].apply(safe_int)
    total_against = df["Total_Against"].apply(safe_int)
    vote_diff = (total_for - total_against).abs()
    winning_bucket = pd.Series("Tie", index=df.index, dtype="object")
    winning_bucket = winning_bucket.mask(total_for > total_against, "For")
    winning_bucket = winning_bucket.mask(total_against > total_for, "Against")
    df["Vote Difference"] = vote_diff
    df["Winning Bucket"] = winning_bucket
    return (
        (vote_diff <= max_vote_diff)
        & winning_bucket.isin(["For", "Against"])
        & (df["Vote Bucket"] == winning_bucket)
    )


def apply_filters_json(
    summary_df: pd.DataFrame,
    *,
    filter_mode: str,
    search_term: str = "",
    bill_id_filters: Optional[List[str]] = None,
    year_selection: Optional[List[int]] = None,
    party_focus_option: str = "Legislator's Party",
    minority_percent: int = 20,
    min_group_votes: int = 0,
    max_vote_diff: int = 5,
    sponsor_metadata: Optional[Dict[Tuple[str, str], Dict[str, object]]] = None,
    selected_legislator: Optional[str] = None,
    legislator_party_label: str = "",
) -> Tuple[pd.DataFrame, int]:
    df = summary_df.copy()

    if year_selection:
        df = df[df["Year"].isin(year_selection)].copy()

    if search_term:
        description_mask = df["Bill Description"].astype(str).str.contains(
            search_term, case=False, na=False
        )
        df = df[description_mask].copy()

    if bill_id_filters:
        normalized_ids = {
            str(bill_id).strip()
            for bill_id in bill_id_filters
            if str(bill_id).strip()
        }
        if normalized_ids:
            df = df[df["Bill ID"].astype(str).str.strip().isin(normalized_ids)].copy()
            if df.empty:
                raise ValueError("No vote records found for the provided Bill IDs.")

    if df.empty:
        raise ValueError("No vote records found for the selected criteria.")

    if filter_mode == "Search By Term" and not search_term:
        raise ValueError("Enter a search term to use the 'Search By Term' vote type.")

    if filter_mode == "Skipped Votes":
        vote_text = df["Vote"].astype(str).str.strip().str.lower()
        skip_mask = ~(
            vote_text.str.startswith("yea")
            | vote_text.str.startswith("nay")
            | vote_text.str.startswith("aye")
        )
        df = df[skip_mask].copy()
        if df.empty:
            raise ValueError("No skipped votes found for the selected criteria.")
    if filter_mode == "Sponsored/Cosponsored Bills":
        sponsor_series = df.get("Sponsorship Status")
        if sponsor_series is None:
            sponsor_series = pd.Series([""] * len(df), index=df.index)
        sponsor_mask = sponsor_series.astype(str).str.strip() != ""
        df = df[sponsor_mask].copy()
        existing_keys: Set[Tuple[str, str]] = {
            (str(session).strip(), str(bill_id).strip())
            for session, bill_id in zip(df.get("Session", []), df.get("Bill ID", []))
        }
        extra_rows: List[Dict[str, object]] = []
        if sponsor_metadata and selected_legislator:
            extra_rows = _create_sponsor_only_rows(
                sponsor_metadata,
                existing_keys,
                selected_legislator,
                legislator_party_label,
            )
        if extra_rows:
            df = pd.concat([df, pd.DataFrame(extra_rows)], ignore_index=True)
        df["Sponsorship Status"] = df["Sponsorship Status"].fillna("").astype(str)
        if df.empty:
            raise ValueError(
                "No sponsored or co-sponsored bills found for the selected legislator."
            )

    df["Roll Call ID"] = pd.to_numeric(df["Roll Call ID"], errors="coerce").astype("Int64")

    def calc_metrics(row: pd.Series):
        bucket = row["Vote Bucket"]
        party = row["Person Party"]
        metrics = {
            "party_bucket_votes": None,
            "party_total_votes": None,
            "party_share": None,
            "chamber_bucket_votes": None,
            "chamber_total_votes": None,
            "chamber_share": None,
        }

        if party:
            party_bucket_col = f"{party}_{bucket}"
            party_total_col = f"{party}_Total"
            party_bucket = safe_int(row.get(party_bucket_col))
            party_total = safe_int(row.get(party_total_col))
            metrics["party_bucket_votes"] = party_bucket
            metrics["party_total_votes"] = party_total
            metrics["party_share"] = party_bucket / party_total if party_total else None

        chamber_bucket = safe_int(row.get(f"Total_{bucket}"))
        chamber_total = safe_int(row.get("Total_Total"))
        metrics["chamber_bucket_votes"] = chamber_bucket
        metrics["chamber_total_votes"] = chamber_total
        metrics["chamber_share"] = chamber_bucket / chamber_total if chamber_total else None
        return pd.Series(metrics)

    metrics_df = df.apply(calc_metrics, axis=1)
    df = pd.concat([df, metrics_df], axis=1)

    df["focus_party_label"] = df["Person Party"]
    df["focus_party_bucket_votes"] = df["party_bucket_votes"]
    df["focus_party_total_votes"] = df["party_total_votes"]
    df["focus_party_share"] = df["party_share"]

    focus_party_key = FOCUS_PARTY_LOOKUP.get(party_focus_option)
    if filter_mode == "Votes Against Party" and focus_party_key:
        focus_display_label = (
            "Independent" if focus_party_key == "Other" else party_focus_option
        )

        def calc_focus_metrics(row: pd.Series):
            bucket = row["Vote Bucket"]
            bucket_votes = safe_int(row.get(f"{focus_party_key}_{bucket}"))
            total_votes = safe_int(row.get(f"{focus_party_key}_Total"))
            share = bucket_votes / total_votes if total_votes else None
            return pd.Series(
                {
                    "focus_party_label": focus_display_label,
                    "focus_party_bucket_votes": bucket_votes,
                    "focus_party_total_votes": total_votes,
                    "focus_party_share": share,
                }
            )

        focus_metrics = df.apply(calc_focus_metrics, axis=1)
        df[
            [
                "focus_party_label",
                "focus_party_bucket_votes",
                "focus_party_total_votes",
                "focus_party_share",
            ]
        ] = focus_metrics

    deciding_condition = None
    if filter_mode == "Deciding Votes":
        deciding_condition = _apply_deciding_vote_filter(df, max_vote_diff)

    apply_party_filter = filter_mode in {"Votes Against Party", "Minority Votes"}
    apply_chamber_filter = filter_mode == "Minority Votes"
    threshold_ratio = (
        minority_percent / 100.0 if (apply_party_filter or apply_chamber_filter) else None
    )
    min_votes = min_group_votes if (apply_party_filter or apply_chamber_filter) else 0

    filters = []
    if apply_party_filter:
        party_condition = (
            df["focus_party_share"].notna()
            & (df["focus_party_total_votes"] >= min_votes)
            & (df["focus_party_share"] <= threshold_ratio)
        )
        filters.append(party_condition)
    if apply_chamber_filter:
        chamber_condition = (
            df["chamber_share"].notna()
            & (df["chamber_total_votes"] >= min_votes)
            & (df["chamber_share"] <= threshold_ratio)
        )
        filters.append(chamber_condition)
    if filter_mode == "Deciding Votes" and deciding_condition is not None:
        filters.append(deciding_condition)

    pre_filter_count = len(df)

    if filters:
        mask = filters[0]
        for condition in filters[1:]:
            mask &= condition
        filtered_df = df[mask].copy()
    else:
        filtered_df = df.copy()

    dedupe_keys = [col for col in ["Roll Call ID", "Person"] if col in filtered_df.columns]
    if dedupe_keys:
        filtered_df = filtered_df.drop_duplicates(subset=dedupe_keys).reset_index(drop=True)

    if filtered_df.empty:
        if filter_mode == "Skipped Votes":
            raise ValueError("No skipped votes found for the selected criteria.")
        if filter_mode == "Votes Against Party":
            raise ValueError(
                "No votes found where the legislator sided with the specified minority."
            )
        if filter_mode == "Minority Votes":
            raise ValueError(
                "No votes found where the legislator and chamber were both in the minority."
            )
        if filter_mode == "Deciding Votes":
            raise ValueError(
                "No votes found within the specified deciding vote margin."
            )
        if filter_mode == "Search By Term":
            raise ValueError(
                "No vote records found matching the provided search term."
            )
        raise ValueError("No vote records found for the selected criteria.")

    return filtered_df, pre_filter_count


def _normalize_bill_marker(value: Optional[str]) -> str:
    if value is None:
        return ""
    return str(value).strip().upper()


def _row_overlap_tokens(row: pd.Series) -> Set[Tuple]:
    tokens: Set[Tuple] = set()
    session_label = _normalize_bill_marker(row.get("Session"))
    roll_call_id = safe_int(row.get("Roll Call ID"))
    if roll_call_id:
        tokens.add(("roll_call", roll_call_id))
    bill_id = _normalize_bill_marker(row.get("Bill ID"))
    if session_label and bill_id:
        tokens.add(("bill_id", session_label, bill_id))
    bill_number = _normalize_bill_marker(row.get("Bill Number"))
    if session_label and bill_number:
        tokens.add(("bill_number", session_label, bill_number))
    crossfile_id = _normalize_bill_marker(row.get("Cross-file Bill ID"))
    if session_label and crossfile_id:
        tokens.add(("crossfile_id", session_label, crossfile_id))
        tokens.add(("bill_id", session_label, crossfile_id))
    crossfile_number = _normalize_bill_marker(row.get("Cross-file Bill Number"))
    if session_label and crossfile_number:
        tokens.add(("crossfile_number", session_label, crossfile_number))
        tokens.add(("bill_number", session_label, crossfile_number))
    return tokens


def _apply_filters_for_package(
    legislator_name: str,
    package: Dict[str, object],
    common_filter_kwargs: Dict[str, object],
) -> Tuple[pd.DataFrame, int, List[Set[Tuple]], Set[Tuple]]:
    df = package.get("df")
    if df is None:
        raise ValueError(f"No vote dataset available for {legislator_name}.")
    params = dict(common_filter_kwargs)
    params["sponsor_metadata"] = package.get("sponsor_metadata")
    params["selected_legislator"] = legislator_name
    params["legislator_party_label"] = package.get("legislator_party_label", "")
    filtered_df, total_count = apply_filters_json(df, **params)
    row_tokens: List[Set[Tuple]] = []
    combined_tokens: Set[Tuple] = set()
    for _, row in filtered_df.iterrows():
        tokens = _row_overlap_tokens(row)
        row_tokens.append(tokens)
        combined_tokens.update(tokens)
    return filtered_df, total_count, row_tokens, combined_tokens


def _filter_with_legislator_overlap(
    primary_name: str,
    multi_legislator_data: Dict[str, Dict[str, object]],
    comparison_legislators: List[str],
    common_filter_kwargs: Dict[str, object],
) -> Tuple[pd.DataFrame, int]:
    primary_package = multi_legislator_data.get(primary_name)
    if primary_package is None:
        raise ValueError("Primary legislator dataset is unavailable. Regenerate the summary.")
    (
        primary_df,
        total_count,
        primary_row_tokens,
        primary_token_set,
    ) = _apply_filters_for_package(
        primary_name,
        primary_package,
        common_filter_kwargs,
    )
    if not comparison_legislators:
        return primary_df, total_count
    overlap_tokens = set(primary_token_set)
    if not overlap_tokens:
        raise ValueError(
            "No qualifying roll-call or cross-file identifiers found for the primary legislator."
        )
    for comp_name in comparison_legislators:
        package = multi_legislator_data.get(comp_name)
        if package is None:
            raise ValueError(f"No cached dataset available for {comp_name}. Regenerate the summary.")
        _, _, _, comp_tokens = _apply_filters_for_package(comp_name, package, common_filter_kwargs)
        overlap_tokens &= comp_tokens
        if not overlap_tokens:
            break
    if not overlap_tokens:
        raise ValueError(
            "No overlapping votes found for the selected legislators with the current filters."
        )
    mask = [
        bool(tokens & overlap_tokens)
        for tokens in primary_row_tokens
    ]
    filtered_subset = primary_df.loc[mask].copy()
    if filtered_subset.empty:
        raise ValueError(
            "Overlapping votes were identified, but none remained after applying all filters."
        )
    return filtered_subset, total_count


def main() -> None:
    st.set_page_config(page_title="LegiScan JSON Vote Explorer", layout="wide")
    st.title("LegiScan JSON Vote Explorer (JSON Beta)")
    st.caption("Load LegiScan JSON archives by state, pick a legislator, and download their vote history.")
    with st.expander("How this workflow works", expanded=False):
        st.markdown(
            """
            1. **Pick state archives** in the sidebar. When remote storage is configured, the app only downloads the ZIPs for the states you select.
            2. **Choose a legislator and view** (All Votes, Votes Against Party, Sponsored Bills, etc.), then click **Generate summary** to cache the dataset.
            3. **Filter & explore** using year, search term, minority/deciding controls, and review the live table to confirm the subset you need.
            4. **Export results** with the download buttons: the current filtered sheet, a Word bullet summary, or the multi-view workbook.
            """
        )

    state_labels, state_codes = _render_state_filter()
    if not state_codes:
        st.info("Select at least one state to continue.")
        st.stop()

    selected_archive_names = _collect_archives_for_states(state_codes)
    if not selected_archive_names:
        readable_states = ", ".join(state_labels)
        st.warning(f"No JSON archives found for: {readable_states}.")
        st.stop()

    archives_snapshot = tuple(sorted(selected_archive_names))

    try:
        selected_paths = _resolve_archives(selected_archive_names)
    except FileNotFoundError as exc:
        st.error(str(exc))
        st.stop()

    paths_snapshot = tuple(str(path) for path in selected_paths)
    with st.spinner("Discovering legislators..."):
        legislator_names, dataset_state = load_legislators_for_archives(paths_snapshot)


    if not legislator_names:
        st.warning("No legislators found in the selected archives.")
        st.stop()

    if state_labels:
        state_display = ", ".join(state_labels)
    else:
        state_display = (
            STATE_CODE_TO_NAME.get(dataset_state, dataset_state)
            if dataset_state
            else ALL_STATES_LABEL
        )
    st.caption(f"Archive selection: {state_display}")

    selected_legislator = st.sidebar.selectbox(
        "Legislator",
        legislator_names,
        index=0,
        key="json_legislator_select",
    )
    comparison_options = [name for name in legislator_names if name != selected_legislator]
    additional_legislators = st.sidebar.multiselect(
        "Additional legislators (overlap)",
        options=comparison_options,
        key="json_additional_legislators",
        help="Limit results to votes shared with these legislators (same roll call and filters).",
    )

    filter_mode = st.sidebar.selectbox(
        "Vote type",
        options=FILTER_OPTIONS,
        index=0,
        key="json_filter_mode",
        help="Choose a predefined view of the legislator's voting record.",
    )
    search_term = st.sidebar.text_input(
        "Search term (bill description)",
        value="",
        key="json_search_term",
    )
    bill_id_filter_text = st.sidebar.text_area(
        "Bill IDs (optional)",
        value="",
        key="json_bill_id_filter",
        help="Comma or newline separated list of Bill IDs to include.",
    )
    bill_id_filters = _parse_bill_id_list(bill_id_filter_text)
    bullet_amendments = st.sidebar.checkbox(
        "Bullet bill amendments?",
        value=False,
        help="Include amendment roll-call votes when building the bullet summary.",
    )

    party_focus_option = "Legislator's Party"
    minority_percent = 20
    min_group_votes = 0
    max_vote_diff = 5

    if filter_mode == "Votes Against Party":
        party_focus_option = st.sidebar.selectbox(
            "Party voting against",
            options=list(FOCUS_PARTY_LOOKUP.keys()),
            index=0,
            key="json_party_focus",
            help="Choose which party's vote breakdown to compare against.",
        )
        minority_percent = st.sidebar.slider(
            "Minority threshold (%)",
            min_value=0,
            max_value=100,
            value=20,
            key="json_votes_against_threshold",
            help="Keep votes where the selected party supported the legislator at or below this percentage.",
        )
        min_group_votes = st.sidebar.slider(
            "Minimum votes in group",
            min_value=0,
            max_value=200,
            value=5,
            key="json_votes_against_min_votes",
            help="Ignore records where the compared party cast fewer total votes than this threshold.",
        )
        st.sidebar.caption("Shows votes where the legislator sided with a minority of the chosen party.")
    elif filter_mode == "Minority Votes":
        minority_percent = st.sidebar.slider(
            "Minority threshold (%)",
            min_value=0,
            max_value=100,
            value=20,
            key="json_minority_threshold",
            help="Keep votes where the legislator's party supported their position at or below this percentage.",
        )
        min_group_votes = st.sidebar.slider(
            "Minimum votes in group",
            min_value=0,
            max_value=200,
            value=5,
            key="json_minority_min_votes",
            help="Ignore records where the compared group cast fewer total votes than this threshold.",
        )
        st.sidebar.caption("Shows votes where the legislator and chamber were both in the minority.")
    elif filter_mode == "Deciding Votes":
        max_vote_diff = st.sidebar.slider(
            "Maximum votes difference",
            min_value=1,
            max_value=50,
            value=5,
            key="json_deciding_max_diff",
            help="Limit to votes where the margin between Yeas and Nays is within this amount.",
        )
        st.sidebar.caption("Shows votes where the legislator's side prevailed by the specified margin or less.")
    elif filter_mode == "Search By Term":
        st.sidebar.caption("Shows votes where the bill description matches the search term.")
    elif filter_mode == "Skipped Votes":
        st.sidebar.caption("Shows votes where the legislator did not cast a Yea or Nay.")
    elif filter_mode == "Sponsored/Cosponsored Bills":
        st.sidebar.caption("Shows votes on bills the legislator sponsored or co-sponsored.")

    generate_summary = st.button("Generate summary", type="primary")

    sponsor_metadata: Dict[Tuple[str, str], Dict[str, object]] = {}
    legislator_party_label = ""

    if generate_summary:
        targets = [selected_legislator] + [
            name for name in additional_legislators if name != selected_legislator
        ]
        seen_targets: Set[str] = set()
        ordered_targets: List[str] = []
        for name in targets:
            if name and name not in seen_targets:
                ordered_targets.append(name)
                seen_targets.add(name)
        multi_legislator_data: Dict[str, Dict[str, object]] = {}
        for name in ordered_targets:
            with st.spinner(f"Compiling votes for {name}..."):
                try:
                    package = _build_legislator_dataset(selected_paths, name)
                except Exception as exc:
                    st.error(f"Failed to build vote summary for {name}: {exc}")
                    st.stop()
            multi_legislator_data[name] = package
        primary_package = multi_legislator_data[selected_legislator]
        rows = primary_package["rows"]
        summary_df = primary_package["df"]
        sponsor_metadata = primary_package.get("sponsor_metadata", {})
        legislator_party_label = primary_package.get("legislator_party_label", "")
        bill_vote_metadata = primary_package.get("bill_vote_metadata", {})
        st.session_state[SESSION_CACHE_KEY] = {
            "rows": rows,
            "df": summary_df,
            "legislator": selected_legislator,
            "archives": archives_snapshot,
            "sponsor_metadata": sponsor_metadata,
            "legislator_party_label": legislator_party_label,
            "bill_vote_metadata": bill_vote_metadata,
            "multi_legislator_data": multi_legislator_data,
            "additional_legislators": additional_legislators,
        }

    cached = st.session_state.get(SESSION_CACHE_KEY)
    if cached:
        if cached.get("archives") != archives_snapshot or cached.get("legislator") not in legislator_names:
            cached = None
        else:
            cached_compare = set(cached.get("additional_legislators", []))
            if cached_compare != set(additional_legislators):
                cached = None

    if not cached:
        st.info("Click **Generate summary** to build the vote dataset.")
        st.stop()

    legislator = cached["legislator"]
    multi_legislator_data: Dict[str, Dict[str, object]] = cached.get("multi_legislator_data") or {}
    if not multi_legislator_data:
        multi_legislator_data = {
            legislator: {
                "rows": cached.get("rows", []),
                "df": cached.get("df"),
                "sponsor_metadata": cached.get("sponsor_metadata", {}),
                "legislator_party_label": cached.get("legislator_party_label", ""),
                "bill_vote_metadata": cached.get("bill_vote_metadata", {}),
            }
        }
    primary_package = multi_legislator_data.get(legislator)
    if primary_package is None:
        primary_package = {
            "rows": cached.get("rows", []),
            "df": cached.get("df"),
            "sponsor_metadata": cached.get("sponsor_metadata", {}),
            "legislator_party_label": cached.get("legislator_party_label", ""),
            "bill_vote_metadata": cached.get("bill_vote_metadata", {}),
        }
        multi_legislator_data[legislator] = primary_package
        cached["multi_legislator_data"] = multi_legislator_data
        st.session_state[SESSION_CACHE_KEY] = cached
    rows = primary_package.get("rows", [])
    summary_df = primary_package.get("df")
    sponsor_metadata = primary_package.get("sponsor_metadata", {})
    legislator_party_label = primary_package.get("legislator_party_label", "")
    bill_vote_metadata = primary_package.get("bill_vote_metadata", {})
    if "Sponsorship Status" not in summary_df.columns:
        summary_df["Sponsorship Status"] = ""

    year_options = sorted(
        int(year)
        for year in summary_df["Year"].dropna().unique().tolist()
        if pd.notna(year)
    )
    st.session_state["json_year_options"] = year_options
    stored_years = st.session_state.get("json_year_selection", [])
    stored_years = [year for year in stored_years if year in year_options]
    if not stored_years and year_options:
        stored_years = year_options
    year_selection = st.sidebar.multiselect(
        "Year",
        options=year_options,
        default=stored_years,
        key="json_year_selection_widget",
        help="Restrict votes to selected calendar years.",
        disabled=not year_options,
    )
    st.session_state["json_year_selection"] = year_selection

    comparison_legislators = [
        name for name in additional_legislators if name in multi_legislator_data
    ]
    common_filter_kwargs = {
        "filter_mode": filter_mode,
        "search_term": search_term,
        "bill_id_filters": bill_id_filters,
        "year_selection": year_selection,
        "party_focus_option": party_focus_option,
        "minority_percent": minority_percent,
        "min_group_votes": min_group_votes,
        "max_vote_diff": max_vote_diff,
    }
    try:
        filtered_df, total_count = _filter_with_legislator_overlap(
            legislator,
            multi_legislator_data,
            comparison_legislators,
            common_filter_kwargs,
        )
    except ValueError as exc:
        st.warning(str(exc))
        st.stop()

    filtered_count = len(filtered_df)
    overlap_note = ""
    if comparison_legislators:
        overlap_note = f" (overlap across {1 + len(comparison_legislators)} legislators)"
    st.success(
        f"Compiled {total_count} votes for {legislator}. "
        f"Showing {filtered_count} after filters{overlap_note}."
    )

    display_df = filtered_df.copy()
    if "Date_dt" in display_df.columns:
        display_df["Date"] = display_df["Date_dt"].dt.date
    display_columns = [
        "Date",
        "Session",
        "Bill Number",
        "Bill Motion",
        "Chamber",
        "Vote",
        "Vote Bucket",
        "Result",
    ]
    st.dataframe(display_df[display_columns], width="stretch", height=400)

    export_df = filtered_df.copy()
    if "Date_dt" in export_df.columns:
        export_df = export_df.drop(columns=["Date_dt"])
    export_df = export_df.reindex(columns=WORKBOOK_HEADERS)
    export_df = export_df.fillna("").infer_objects(copy=False)
    excel_rows = export_df.values.tolist()
    excel_buffer = io.BytesIO()
    write_workbook(excel_rows, legislator, excel_buffer)
    st.download_button(
        label="Download filtered Excel sheet",
        data=excel_buffer.getvalue(),
        file_name=f"{legislator.replace(' ', '_')}_JSON_Votes.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    if not filtered_df.empty:
        bullet_buffer = _build_json_bullet_summary_doc(
            filtered_df,
            legislator,
            filter_mode,
            state_display,
            bill_vote_metadata=bill_vote_metadata,
            full_summary_df=summary_df,
            include_amendments=bullet_amendments,
        )
        st.download_button(
            label="Download bullet summary",
            data=bullet_buffer.getvalue(),
            file_name=f"{legislator.replace(' ', '_')}_{filter_mode.replace('/', '_').replace(' ', '_')}_summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="json_bullet_summary_download",
        )
    generate_workbook_clicked = False
    if not filtered_df.empty:
        generate_workbook_clicked = st.button(
            "Generate all views workbook",
            key="json_full_workbook_button",
            help="Compile each predefined view into a single Excel workbook.",
        )

    if generate_workbook_clicked:
        workbook_views: List[Tuple[str, List[str], List[List]]] = []
        empty_views: List[str] = []
        stored_party_focus = st.session_state.get("json_party_focus", party_focus_option)
        stored_votes_against_threshold = st.session_state.get("json_votes_against_threshold", 20)
        stored_votes_against_min_votes = st.session_state.get("json_votes_against_min_votes", 5)
        stored_minority_threshold = st.session_state.get("json_minority_threshold", 20)
        stored_minority_min_votes = st.session_state.get("json_minority_min_votes", 5)
        stored_deciding_max_diff = st.session_state.get("json_deciding_max_diff", 5)

        base_params = {
            "search_term": "",
            "bill_id_filters": bill_id_filters,
            "year_selection": year_selection or None,
            "party_focus_option": "Legislator's Party",
            "minority_percent": 20,
            "min_group_votes": 0,
            "max_vote_diff": 5,
        }

        for view_name in WORKBOOK_VIEWS:
            params = base_params.copy()
            if view_name == "Votes Against Party":
                params.update(
                    {
                        "party_focus_option": stored_party_focus,
                        "minority_percent": stored_votes_against_threshold,
                        "min_group_votes": stored_votes_against_min_votes,
                    }
                )
            elif view_name == "Minority Votes":
                params.update(
                    {
                        "minority_percent": stored_minority_threshold,
                        "min_group_votes": stored_minority_min_votes,
                    }
                )
            elif view_name == "Deciding Votes":
                params.update({"max_vote_diff": stored_deciding_max_diff})
            try:
                sheet_df, _ = _filter_with_legislator_overlap(
                    legislator,
                    multi_legislator_data,
                    comparison_legislators,
                    {
                        "filter_mode": view_name,
                        **params,
                    },
                )
            except ValueError:
                empty_views.append(view_name)
                continue
            export_sheet = sheet_df.copy()
            export_sheet = export_sheet.drop(columns=["Date_dt"], errors="ignore")
            export_sheet = export_sheet.reindex(columns=WORKBOOK_HEADERS)
            export_sheet = export_sheet.fillna("").infer_objects(copy=False)
            sheet_rows = export_sheet.values.tolist()
            workbook_views.append((view_name, WORKBOOK_HEADERS, sheet_rows))

        if not workbook_views:
            st.warning("No data available for the selected workbook views.")
        else:
            workbook_buffer = io.BytesIO()
            write_multi_sheet_workbook(workbook_views, workbook_buffer)
            workbook_buffer.seek(0)
            st.download_button(
                label="Download vote summary workbook",
                data=workbook_buffer.getvalue(),
                file_name=f"{legislator.replace(' ', '_')}_JSON_full_workbook.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="json_full_workbook_download",
            )
            if empty_views:
                st.info("No data available for: " + ", ".join(empty_views))


if __name__ == "__main__":
    main()
