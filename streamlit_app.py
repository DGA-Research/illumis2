import base64
import csv
import datetime as dt
import hashlib
import io
import re
import tempfile
import zipfile
from pathlib import Path
from functools import lru_cache
from typing import Dict, Iterable, List, Optional, Tuple, Set, Union

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt
from openpyxl import Workbook

import pandas as pd
import requests
import streamlit as st

from generate_kristin_robbins_votes import (
    WORKBOOK_HEADERS,
    collect_legislator_names,
    collect_vote_rows,
    determine_dataset_state,
    gather_session_csv_dirs,
    collect_person_vote_map,
)

LOCAL_ARCHIVE_DIR = Path(__file__).resolve().parent / "bulkLegiData"
BUNDLED_ARCHIVE_SESSION_KEY = "bundled_archive_selection"
VIEW_STATE_SESSION_KEY = "current_vote_view_state"
ALL_STATES_LABEL = "All States"
STATE_CHOICES = [
    ("Alabama", "AL"),
    ("Alaska", "AK"),
    ("Arizona", "AZ"),
    ("Arkansas", "AR"),
    ("California", "CA"),
    ("Colorado", "CO"),
    ("Connecticut", "CT"),
    ("Delaware", "DE"),
    ("Florida", "FL"),
    ("Georgia", "GA"),
    ("Hawaii", "HI"),
    ("Idaho", "ID"),
    ("Illinois", "IL"),
    ("Indiana", "IN"),
    ("Iowa", "IA"),
    ("Kansas", "KS"),
    ("Kentucky", "KY"),
    ("Louisiana", "LA"),
    ("Maine", "ME"),
    ("Maryland", "MD"),
    ("Massachusetts", "MA"),
    ("Michigan", "MI"),
    ("Minnesota", "MN"),
    ("Mississippi", "MS"),
    ("Missouri", "MO"),
    ("Montana", "MT"),
    ("Nebraska", "NE"),
    ("Nevada", "NV"),
    ("New Hampshire", "NH"),
    ("New Jersey", "NJ"),
    ("New Mexico", "NM"),
    ("New York", "NY"),
    ("North Carolina", "NC"),
    ("North Dakota", "ND"),
    ("Ohio", "OH"),
    ("Oklahoma", "OK"),
    ("Oregon", "OR"),
    ("Pennsylvania", "PA"),
    ("Rhode Island", "RI"),
    ("South Carolina", "SC"),
    ("South Dakota", "SD"),
    ("Tennessee", "TN"),
    ("Texas", "TX"),
    ("Utah", "UT"),
    ("Vermont", "VT"),
    ("Virginia", "VA"),
    ("Washington", "WA"),
    ("West Virginia", "WV"),
    ("Wisconsin", "WI"),
    ("Wyoming", "WY"),
]
STATE_NAME_TO_CODE = {name: code for name, code in STATE_CHOICES}
PARTY_DISPLAY_MAP = {
    "Democrat": "Democrat",
    "Republican": "Republican",
    "Other": "Independent",
}
FOCUS_PARTY_LOOKUP = {
    "Legislator's Party": None,
    "Democrat": "Democrat",
    "Republican": "Republican",
    "Independent": "Other",
}
SPONSOR_DROP_COLUMNS = [
    "Roll Details",
    "Roll Call ID",
    "Vote",
    "Vote Bucket",
    "Bill ID",
    "Result",
    "Democrat_For",
    "Democrat_Against",
    "Democrat_Absent",
    "Democrat_Not",
    "Democrat_Total",
    "Republican_For",
    "Republican_Against",
    "Republican_Absent",
    "Republican_Not",
    "Republican_Total",
    "Other_For",
    "Other_Against",
    "Other_Absent",
    "Other_Not",
    "Other_Total",
    "Total_For",
    "Total_Against",
    "Total_Absent",
    "Total_Not",
    "Total_Total",
]

ARCHIVE_INFO_PATTERN = re.compile(
    r"^(?P<state>[A-Z]{2})_(?P<start>\d{4})-(?P<end>\d{4})_(?P<label>.+?)_CSV_(?P<stamp>\d{8})(?:_[0-9a-f]+)?\.zip$",
    re.IGNORECASE,
)

EXTRACTED_ARCHIVE_CACHE_KEY = "_cached_extracted_archives"
_LOCAL_EXTRACTION_CACHE: Dict[str, Dict[str, object]] = {}


class ArchiveUpToDate(Exception):
    def __init__(self, filename: str):
        self.filename = filename
        super().__init__(f"{filename} is already up to date.")

@lru_cache(maxsize=512)
def _parse_archive_session_info(name: str) -> Optional[Dict[str, object]]:
    match = ARCHIVE_INFO_PATTERN.match(name.strip())
    if not match:
        return None
    start_year = int(match.group("start"))
    end_year = int(match.group("end"))
    label = match.group("label")
    stamp_str = match.group("stamp")
    try:
        stamp_date = dt.datetime.strptime(stamp_str, "%Y%m%d").date()
    except ValueError:
        stamp_date = None
    return {
        "state": match.group("state").upper(),
        "start_year": start_year,
        "end_year": end_year,
        "label": label,
        "stamp_date": stamp_date,
    }


def _archive_session_status(name: str, *, reference_date: Optional[dt.date] = None) -> str:
    info = _parse_archive_session_info(name)
    if not info:
        return "unknown"
    ref_date = reference_date or dt.date.today()
    start_date = dt.date(info["start_year"], 1, 1)
    end_date = dt.date(info["end_year"], 12, 31)
    if ref_date < start_date:
        return "upcoming"
    if ref_date > end_date:
        return "past"
    return "current"


def _format_archive_option(name: str) -> str:
    info = _parse_archive_session_info(name)
    if not info:
        return name
    status = _archive_session_status(name)
    status_label = {
        "current": "CURRENT",
        "upcoming": "UPCOMING",
        "past": "PAST",
        "unknown": "UNKNOWN",
    }[status]
    descriptor = info["label"].replace("_", " ")
    parts = [status_label, descriptor]
    if info.get("stamp_date"):
        parts.append(f"updated {info['stamp_date'].isoformat()}")
    return f"{name} ({'; '.join(parts)})"


def _group_archives_by_status(
    names: Iterable[str], *, reference_date: Optional[dt.date] = None
) -> Dict[str, List[str]]:
    ref_date = reference_date or dt.date.today()
    grouped: Dict[str, List[str]] = {
        "current": [],
        "upcoming": [],
        "past": [],
        "unknown": [],
    }
    for name in sorted(set(name for name in names if name)):
        status = _archive_session_status(name, reference_date=ref_date)
        grouped.setdefault(status, []).append(name)
    return grouped


def _latest_last_action_from_payload(payload: bytes) -> Optional[dt.date]:
    try:
        return _collect_latest_action_date([payload])
    except zipfile.BadZipFile:
        return None


def _infer_payload_freshness(name: str, payload: bytes) -> Optional[dt.date]:
    freshest = _latest_last_action_from_payload(payload)
    if freshest:
        return freshest
    info = _parse_archive_session_info(name)
    if info:
        stamp_date = info.get("stamp_date")
        if isinstance(stamp_date, dt.date):
            return stamp_date
    return None


def _compose_row_key(row: pd.Series) -> str:
    person = str(row.get("Person") or "").strip()
    rcid = row.get("Roll Call ID")
    if pd.notna(rcid):
        try:
            numeric_id = int(rcid)
            base_key = f"RCID_{numeric_id}"
        except (TypeError, ValueError):
            base_key = f"RCID_{str(rcid).strip()}"
        return f"{base_key}::{person}" if person else base_key
    session = str(row.get("Session") or "").strip()
    bill_number = str(row.get("Bill Number") or "").strip()
    vote = str(row.get("Vote") or "").strip()
    chamber = str(row.get("Chamber") or "").strip()
    date_value = ""
    date_raw = row.get("Date_dt")
    if pd.notna(date_raw):
        if isinstance(date_raw, dt.datetime):
            date_value = date_raw.strftime("%Y-%m-%d")
        else:
            date_value = str(date_raw)
    else:
        date_value = str(row.get("Date") or "").strip()
    base_key = f"ALT_{session}_{bill_number}_{chamber}_{vote}_{date_value}"
    return f"{base_key}::{person}" if person else base_key



def _collect_legislators_from_zips(zip_payloads: List[bytes]):
    base_dirs = _get_payload_base_dirs(zip_payloads)
    state = determine_dataset_state(base_dirs)
    names = collect_legislator_names(base_dirs)
    return state, names


def _collect_rows_from_zips(
    zip_payloads: List[bytes], legislator_name: Union[str, Iterable[str]]
) -> Tuple[List[List], List[str]]:
    base_dirs = _get_payload_base_dirs(zip_payloads)
    if isinstance(legislator_name, str) or legislator_name is None:
        target_names = [legislator_name] if legislator_name else []
    else:
        target_names = [name for name in legislator_name if name]
    if not target_names:
        raise ValueError("No legislator selected.")
    missing_legislators: List[str] = []
    all_rows: List[List] = []
    for idx, name in enumerate(target_names):
        try:
            all_rows.extend(collect_vote_rows(base_dirs, name))
        except ValueError:
            missing_legislators.append(name)
            if idx == 0:
                # Primary selection missing data should still halt processing.
                raise
    if not all_rows:
        raise ValueError("No vote records found for the selected legislators.")
    return all_rows, missing_legislators


def _collect_person_votes_from_zips(zip_payloads: List[bytes], legislator_name: str):
    base_dirs = _get_payload_base_dirs(zip_payloads)
    return collect_person_vote_map(base_dirs, legislator_name)


def _collect_years_from_zips(zip_payloads: List[bytes]):
    years = set()
    base_dirs = _get_payload_base_dirs(zip_payloads)
    try:
        csv_dirs = gather_session_csv_dirs(base_dirs)
    except FileNotFoundError:
        return []
    for csv_dir in csv_dirs:
        rollcalls_path = csv_dir / "rollcalls.csv"
        if not rollcalls_path.exists():
            continue
        with rollcalls_path.open(encoding="utf-8") as fh:
            reader = csv.DictReader(fh)
            for row in reader:
                date_str = (row.get("date") or "").strip()
                if not date_str:
                    continue
                try:
                    year = dt.datetime.strptime(date_str, "%Y-%m-%d").year
                except ValueError:
                    continue
                years.add(year)
    return sorted(years)


def _format_type_label(label: str) -> str:
    tokens = re.findall(r"[A-Za-z0-9]+", label or "")
    if not tokens:
        return "ALL"
    return "_".join(tokens).upper()


def _format_name_segment(name: str) -> str:
    tokens = re.findall(r"[A-Za-z0-9]+", name or "")
    if not tokens:
        return "Legislator"
    return "".join(token.capitalize() for token in tokens)


def _normalize_state_segment(state_value: Optional[str]) -> str:
    tokens = re.findall(r"[A-Za-z0-9]+", (state_value or "").upper())
    if not tokens:
        return "UNK"
    return "".join(tokens)


def _make_download_filename(
    legislator_name: str,
    type_label: str,
    *,
    dataset_state: Optional[str] = None,
    fallback_state: Optional[str] = None,
) -> str:
    state_segment = _normalize_state_segment(dataset_state or fallback_state)
    name_segment = _format_name_segment(legislator_name)
    type_segment = _format_type_label(type_label)
    return f"{state_segment}_{name_segment}_{type_segment}.xlsx"


def _make_docx_filename(
    legislator_name: str,
    type_label: str,
    *,
    dataset_state: Optional[str] = None,
    fallback_state: Optional[str] = None,
) -> str:
    base = _make_download_filename(
        legislator_name,
        type_label,
        dataset_state=dataset_state,
        fallback_state=fallback_state,
    )
    if base.lower().endswith(".xlsx"):
        return base[:-5] + ".docx"
    return base + ".docx"


def _list_local_archives() -> List[Path]:
    if not LOCAL_ARCHIVE_DIR.exists():
        return []
    return sorted(
        path for path in LOCAL_ARCHIVE_DIR.glob("*.zip") if path.is_file()
    )


def _archive_matches_state(name: str, state_code: str) -> bool:
    if not state_code:
        return True
    prefix = state_code.upper()
    normalized = name.upper()
    return normalized.startswith(prefix)


def _archive_key(name: str) -> str:
    return name.strip().lower()


def _get_payload_base_dirs(zip_payloads: List[bytes]) -> List[Path]:
    if not zip_payloads:
        return []
    try:
        cache: Dict[str, Dict[str, object]] = st.session_state.setdefault(
            EXTRACTED_ARCHIVE_CACHE_KEY, {}
        )
    except Exception:
        cache = _LOCAL_EXTRACTION_CACHE.setdefault(EXTRACTED_ARCHIVE_CACHE_KEY, {})
    base_dirs: List[Path] = []
    for payload in zip_payloads:
        digest = hashlib.sha256(payload).hexdigest()
        entry = cache.get(digest)
        path_obj: Optional[Path] = None
        if entry:
            path_candidate = entry.get("path")
            if isinstance(path_candidate, Path) and path_candidate.exists():
                path_obj = path_candidate
        if path_obj is None:
            temp_dir = tempfile.TemporaryDirectory()
            with zipfile.ZipFile(io.BytesIO(payload)) as zf:
                zf.extractall(temp_dir.name)
            path_obj = Path(temp_dir.name)
            cache[digest] = {"path": path_obj, "temp_ref": temp_dir}
        base_dirs.append(path_obj)
    return base_dirs


FORBIDDEN_SHEET_CHARS = set('[]:*?/\\')
ARCHIVE_NAME_PATTERN = re.compile(r"^[A-Z]{2}_[A-Za-z0-9_.\-]+\.zip$")
REQUIRED_ARCHIVE_FILES = ("people.csv", "bills.csv", "rollcalls.csv", "votes.csv")
PARTY_CODE_MAP = {
    "D": "Democrat",
    "DEM": "Democrat",
    "DFL": "Democrat",
    "R": "Republican",
    "REP": "Republican",
    "GOP": "Republican",
}
HOUSE_PREFIXES = ("HOUSE", "HJR", "HCR", "HB", "HR", "HC", "HJ", "HS", "H")
SENATE_PREFIXES = ("SENATE", "SJR", "SCR", "SB", "SR", "SC", "SJ", "SS", "S")


def _normalize_party_label(party_code: str) -> str:
    code = (party_code or "").strip().upper()
    if not code:
        return ""
    if code in PARTY_CODE_MAP:
        return PARTY_CODE_MAP[code]
    if code in {"I", "IND", "IND.", "INDP", "INDEPENDENT"}:
        return "Other"
    return "Other"


def _infer_chamber_from_bill(bill_number: str) -> str:
    token = (bill_number or "").strip().upper()
    if not token:
        return ""
    for prefix in SENATE_PREFIXES:
        if token.startswith(prefix):
            return "Senate"
    for prefix in HOUSE_PREFIXES:
        if token.startswith(prefix):
            return "House"
    return ""


def _format_roll_details(roll_row: dict) -> str:
    if not roll_row:
        return ""
    description = (roll_row.get("description") or "").strip()

    def _normalized_count(value):
        try:
            return int(value)
        except (TypeError, ValueError):
            return 0

    yea = _normalized_count(roll_row.get("yea"))
    nay = _normalized_count(roll_row.get("nay"))
    suffix = ""
    if yea or nay:
        suffix = f" ({yea}-Y {nay}-N)"
    if description:
        return f"{description}{suffix}"
    return suffix.strip()


def _build_sponsor_metadata(
    bill_row: dict, roll_list: List[dict], sponsorship_status: str
) -> dict:
    session_id = str(bill_row.get("session_id", "")).strip()
    bill_number = str(bill_row.get("bill_number", "")).strip()
    title = (bill_row.get("title") or "").strip()
    description = (bill_row.get("description") or "").strip()
    url = (bill_row.get("state_link") or bill_row.get("url") or "").strip()
    status_code = (bill_row.get("status") or "").strip()
    status_desc = (bill_row.get("status_desc") or "").strip()
    status_date = (bill_row.get("status_date") or "").strip()
    last_action_date = (bill_row.get("last_action_date") or "").strip()
    last_action = (bill_row.get("last_action") or "").strip()
    chamber = _infer_chamber_from_bill(bill_number)
    roll_call_id = ""
    roll_details = ""
    roll_date = ""
    if roll_list:
        first_roll = roll_list[0]
        roll_call_id = str(first_roll.get("roll_call_id", "")).strip()
        roll_details = _format_roll_details(first_roll)
        roll_date = (first_roll.get("date") or "").strip()
        if not chamber:
            chamber_value = (first_roll.get("chamber") or "").strip()
            chamber = chamber_value.title() if chamber_value else ""
    if not roll_details:
        if last_action:
            roll_details = f"No roll call recorded - {last_action}"
        else:
            roll_details = "No roll call recorded"
    date_str = roll_date or status_date or last_action_date
    display_date = ""
    if date_str:
        try:
            display_date = dt.datetime.strptime(date_str, "%Y-%m-%d").strftime("%m/%d/%Y")
        except ValueError:
            display_date = date_str
    result = 1 if status_code == "4" or status_desc.lower() == "passed" else 0
    bill_motion = title or description or bill_number
    return {
        "bill_id": str(bill_row.get("bill_id", "")).strip(),
        "session_id": session_id,
        "bill_number": bill_number,
        "bill_title": title,
        "bill_description": description,
        "bill_motion": bill_motion,
        "bill_url": url,
        "status_code": status_code,
        "status_desc": status_desc,
        "status_date": status_date,
        "last_action": last_action,
        "last_action_date": last_action_date,
        "roll_call_id": roll_call_id,
        "roll_details": roll_details,
        "roll_date": date_str,
        "excel_date": display_date,
        "result": result,
        "chamber": chamber,
        "sponsorship_status": sponsorship_status,
    }


def _create_sponsor_only_rows(
    sponsor_metadata: dict,
    existing_keys: Set[Tuple[str, str]],
    legislator_name: str,
    legislator_party_label: str,
) -> List[dict]:
    rows: List[dict] = []
    count_start_idx = WORKBOOK_HEADERS.index("Democrat_For")
    party_label = legislator_party_label or "Other"
    for idx, (key, meta) in enumerate(sponsor_metadata.items()):
        if key in existing_keys:
            continue
        roll_call_value = meta.get("roll_call_id", "")
        normalized_roll_id: Optional[int]
        try:
            normalized_roll_id = int(str(roll_call_value))
        except (TypeError, ValueError):
            normalized_roll_id = None
        if normalized_roll_id is None:
            bill_id_raw = meta.get("bill_id", "")
            try:
                normalized_roll_id = -abs(int(str(bill_id_raw)))
            except (TypeError, ValueError):
                normalized_roll_id = -(10**9 + idx)
        row = {header: "" for header in WORKBOOK_HEADERS}
        for header in WORKBOOK_HEADERS[count_start_idx:]:
            row[header] = 0
        row.update(
            {
                "Chamber": meta.get("chamber", ""),
                "Session": meta.get("session_id", ""),
                "Bill Number": meta.get("bill_number", ""),
                "Bill ID": meta.get("bill_id", "")
                or f"{meta.get('session_id', '')}-{meta.get('bill_number', '')}".strip("-"),
                "Bill Motion": meta.get("bill_motion", "") or meta.get("bill_title", ""),
                "URL": meta.get("bill_url", ""),
                "Bill Title": meta.get("bill_title", ""),
                "Bill Description": meta.get("bill_description", "") or meta.get("bill_title", ""),
                "Roll Details": meta.get("roll_details", ""),
                "Roll Call ID": normalized_roll_id,
                "Last Action Date": meta.get("last_action_date", ""),
                "Last Action": meta.get("last_action", ""),
                "Status": meta.get("status_code", ""),
                "Status Description": meta.get("status_desc", ""),
                "Status Date": meta.get("status_date", ""),
                "Person": legislator_name,
                "Person Party": party_label,
                "Date": meta.get("excel_date", ""),
            }
        )
        rows.append(
            {
                **row,
                "Sponsorship Status": meta.get("sponsorship_status", ""),
            }
        )
    return rows


def _collect_bill_metadata(zip_payloads: List[bytes]) -> Dict[Tuple[str, str], Dict[str, str]]:
    metadata: Dict[Tuple[str, str], Dict[str, str]] = {}
    base_dirs = _get_payload_base_dirs(zip_payloads)
    try:
        csv_dirs = gather_session_csv_dirs(base_dirs)
    except FileNotFoundError:
        return metadata
    for csv_dir in csv_dirs:
        bill_lookup: Dict[int, Tuple[str, str]] = {}
        bills_path = csv_dir / "bills.csv"
        if bills_path.exists():
            with bills_path.open(encoding="utf-8") as fh:
                reader = csv.DictReader(fh)
                for row in reader:
                    session_id = str(row.get("session_id", "")).strip()
                    bill_number = str(row.get("bill_number", "")).strip()
                    if not session_id or not bill_number:
                        continue
                    key = (session_id, bill_number)
                    try:
                        bill_id = int(row.get("bill_id", ""))
                    except (TypeError, ValueError):
                        bill_id = None

                    entry = metadata.get(key)
                    if entry is None:
                        entry = {
                            "last_action": (row.get("last_action") or "").strip(),
                            "last_action_date": (row.get("last_action_date") or "").strip(),
                            "status_desc": (row.get("status_desc") or "").strip(),
                            "status_date": (row.get("status_date") or "").strip(),
                            "status_code": (row.get("status") or "").strip(),
                            "title": (row.get("title") or "").strip(),
                            "bill_id": bill_id,
                            "latest_votes": {},
                        }
                        metadata[key] = entry
                    else:
                        entry.update(
                            {
                                "last_action": (row.get("last_action") or "").strip(),
                                "last_action_date": (row.get("last_action_date") or "").strip(),
                                "status_desc": (row.get("status_desc") or "").strip(),
                                "status_date": (row.get("status_date") or "").strip(),
                                "status_code": (row.get("status") or "").strip(),
                                "title": (row.get("title") or "").strip(),
                            }
                        )
                        if bill_id is not None:
                            entry["bill_id"] = bill_id
                        entry.setdefault("latest_votes", {})
                    if bill_id is not None:
                        bill_lookup[bill_id] = key

        rollcalls_path = csv_dir / "rollcalls.csv"
        if rollcalls_path.exists() and bill_lookup:
            with rollcalls_path.open(encoding="utf-8") as fh:
                reader = csv.DictReader(fh)
                for row in reader:
                    try:
                        bill_id = int(row.get("bill_id", ""))
                    except (TypeError, ValueError):
                        continue
                    key = bill_lookup.get(bill_id)
                    if not key:
                        continue
                    chamber_value = (row.get("chamber") or "").strip().title()
                    if chamber_value not in {"Senate", "House"}:
                        continue
                    date_str = (row.get("date") or "").strip()
                    parsed_date: Optional[dt.datetime] = None
                    if date_str:
                        try:
                            parsed_date = dt.datetime.strptime(date_str, "%Y-%m-%d")
                        except ValueError:
                            parsed_date = None
                    entry = metadata.get(key)
                    if entry is None:
                        continue
                    counts = entry.setdefault("latest_votes", {})
                    existing = counts.get(chamber_value)
                    update_counts = False
                    if existing is None:
                        update_counts = True
                    else:
                        existing_date = existing.get("date")
                        if parsed_date and (existing_date is None or parsed_date >= existing_date):
                            update_counts = True
                        elif parsed_date is None and existing_date is None:
                            update_counts = True
                    if update_counts:
                        counts[chamber_value] = {
                            "yea": safe_int(row.get("yea")),
                            "nay": safe_int(row.get("nay")),
                            "date": parsed_date,
                        }
    return metadata


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    if not url:
        paragraph.add_run(text)
        return
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _clean_sentence_fragment(text: str) -> str:
    cleaned = (text or "").strip()
    if cleaned.endswith("."):
        cleaned = cleaned[:-1].strip()
    return cleaned


def _resolve_vote_phrases(vote_bucket: str) -> Tuple[str, str]:
    bucket = (vote_bucket or "").strip().lower()
    if bucket == "for":
        return "VOTED FOR", "voted for"
    if bucket == "against":
        return "VOTED AGAINST", "voted against"
    if bucket == "absent":
        return "WAS ABSENT FOR", "was absent for"
    return "DID NOT VOTE ON", "did not vote on"


def _format_vote_ratio(counts: Optional[Dict[str, object]]) -> str:
    if not counts:
        return ""
    yea = counts.get("yea")
    nay = counts.get("nay")
    if yea is None and nay is None:
        return ""
    yea_str = "?" if yea is None else str(yea)
    nay_str = "?" if nay is None else str(nay)
    return f"{yea_str}-{nay_str}"


def _format_docx_date(date_value: str) -> str:
    text = (date_value or "").strip()
    if not text:
        return text
    date_formats = ("%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y")
    for fmt in date_formats:
        try:
            parsed = dt.datetime.strptime(text, fmt)
            return f"{parsed.month}/{parsed.day}/{str(parsed.year)[-2:]}"
        except ValueError:
            continue
    try:
        parsed = dt.datetime.fromisoformat(text)
        return f"{parsed.month}/{parsed.day}/{str(parsed.year)[-2:]}"
    except ValueError:
        return text


def _compose_status_sentence(
    status_code: str,
    bill_number: str,
    chamber: str,
    last_action: str,
    latest_counts: Optional[Dict[str, Dict[str, object]]] = None,
) -> str:
    status = (status_code or "").strip()
    chamber_text = chamber or "the chamber"
    action_text = _clean_sentence_fragment(last_action)
    bill_ref = bill_number or "the bill"
    bill_ref_upper = (bill_number or "").upper()
    last_action_lower = (last_action or "").lower()
    last_action_upper = (last_action or "").upper()

    if status == "1":
        if action_text:
            return f"{bill_ref} introduced in {chamber_text} and {action_text}."
        return f"{bill_ref} introduced in {chamber_text}."
    if status == "2":
        latest_counts = latest_counts or {}
        normalized_chamber = chamber_text.title()
        chamber_display = normalized_chamber if normalized_chamber in {"House", "Senate"} else chamber_text
        ratio = _format_vote_ratio(latest_counts.get(normalized_chamber))
        if ratio:
            if action_text:
                return f"{bill_ref} passed in {chamber_display} {ratio} and {action_text}."
            return f"{bill_ref} passed in {chamber_display} {ratio}."
        if action_text:
            return f"{bill_ref} passed in {chamber_text} and {action_text}."
        return f"{bill_ref} passed in {chamber_text}."
    if status == "4":
        latest_counts = latest_counts or {}
        senate_ratio = _format_vote_ratio(latest_counts.get("Senate"))
        house_ratio = _format_vote_ratio(latest_counts.get("House"))

        if any(token in last_action_upper for token in ("SB", "HB", "CHAPTER")):
            suffix = f", {action_text}" if action_text else ""
            if senate_ratio and house_ratio:
                return (
                    f"{bill_ref} passed in Senate {senate_ratio} and House {house_ratio}, "
                    f"and signed by governor{suffix}."
                )
            return (
                f"{bill_ref} passed in Senate and House, and signed by governor{suffix}."
            )

        if any(token in bill_ref_upper for token in ("CR", "CM")):
            if senate_ratio and house_ratio:
                if action_text:
                    return (
                        f"{bill_ref} passed in Senate {senate_ratio} and House {house_ratio}, "
                        f"and {action_text}."
                    )
                return (
                    f"{bill_ref} passed in Senate {senate_ratio} and House {house_ratio}."
                )
            if action_text:
                return f"{bill_ref} passed in Senate and House and {action_text}."
            return f"{bill_ref} passed in Senate and House."

        normalized_chamber = chamber_text.title()
        chamber_display = (
            normalized_chamber if normalized_chamber in {"House", "Senate"} else chamber_text
        )
        chamber_ratio = _format_vote_ratio(latest_counts.get(normalized_chamber))
        if chamber_ratio:
            if action_text:
                return f"{bill_ref} passed in {chamber_display} {chamber_ratio} and {action_text}."
            return f"{bill_ref} passed in {chamber_display} {chamber_ratio}."
        if action_text:
            return f"{bill_ref} passed in {chamber_display} and {action_text}."
        return f"{bill_ref} passed in {chamber_display}."
    if status == "5":
        latest_counts = latest_counts or {}
        senate_ratio = _format_vote_ratio(latest_counts.get("Senate"))
        house_ratio = _format_vote_ratio(latest_counts.get("House"))
        primary_first = "house" if chamber_text.lower() == "house" else "senate"
        if primary_first == "house":
            ordered = [("House", house_ratio), ("Senate", senate_ratio)]
        else:
            ordered = [("Senate", senate_ratio), ("House", house_ratio)]
        first_label, first_ratio = ordered[0]
        second_label, second_ratio = ordered[1]
        if first_ratio and second_ratio:
            return (
                f"{bill_ref} passed in {first_label} {first_ratio} "
                f"and {second_label} {second_ratio}, vetoed by governor."
            )
        if first_ratio:
            return f"{bill_ref} passed in {first_label} {first_ratio}, vetoed by governor."
        if second_ratio:
            return f"{bill_ref} passed in {second_label} {second_ratio}, vetoed by governor."
        return f"{bill_ref} passed in Senate and House, vetoed by governor."
    if status == "6":
        if "S" in bill_ref_upper and "house" in last_action_lower:
            if action_text:
                return f"{bill_ref} passed in Senate, {action_text}."
            return f"{bill_ref} passed in Senate."
        if "H" in bill_ref_upper and "senate" in last_action_lower:
            if action_text:
                return f"{bill_ref} passed in House, {action_text}."
            return f"{bill_ref} passed in House."
        if action_text:
            return f"{bill_ref} introduced in {chamber_text}, {action_text}."
        return f"{bill_ref} introduced in {chamber_text}."

    if action_text:
        return f"{bill_ref} {action_text}."
    return f"{bill_ref} status information unavailable."


def _prepare_sponsor_dataframe(df: pd.DataFrame, *, drop_date: bool = True) -> pd.DataFrame:
    working = df.copy()
    if "Bill Number" not in working.columns:
        return working

    if "Date_dt" not in working.columns:
        working["Date_dt"] = pd.to_datetime(working.get("Date"), errors="coerce")

    sort_columns: List[str] = []
    ascending: List[bool] = []
    bill_id_series = working["Bill ID"] if "Bill ID" in working.columns else None
    has_bill_id = False
    if bill_id_series is not None:
        has_bill_id = bill_id_series.astype(str).str.strip().ne("").any()
    if has_bill_id:
        sort_columns.append("Bill ID")
        ascending.append(True)
    else:
        sort_columns.append("Bill Number")
        ascending.append(True)
        if "Session" in working.columns:
            sort_columns.append("Session")
            ascending.append(True)
    sort_columns.append("Date_dt")
    ascending.append(False)
    working = working.sort_values(sort_columns, ascending=ascending)

    if has_bill_id:
        dedupe_subset = ["Bill ID"]
    else:
        dedupe_candidates = [col for col in ("Session", "Bill Number") if col in working.columns]
        if not dedupe_candidates and "Roll Call ID" in working.columns:
            dedupe_subset = ["Roll Call ID"]
        else:
            dedupe_subset = dedupe_candidates or ["Bill Number"]
    working = working.drop_duplicates(subset=dedupe_subset, keep="first").reset_index(drop=True)

    working = working.drop(columns=SPONSOR_DROP_COLUMNS, errors="ignore")
    if drop_date:
        working = working.drop(columns=["Date_dt"], errors="ignore")
    return working


def _sanitize_sheet_title(title: str, used_titles: Set[str]) -> str:
    cleaned = "".join("_" if ch in FORBIDDEN_SHEET_CHARS else ch for ch in title)
    cleaned = cleaned.strip() or "Sheet"
    cleaned = cleaned[:31]
    base = cleaned
    counter = 1
    while cleaned in used_titles:
        suffix = f"_{counter}"
        cleaned = (base[: 31 - len(suffix)] + suffix) if len(base) + len(suffix) > 31 else base + suffix
        counter += 1
    used_titles.add(cleaned)
    return cleaned


def _write_single_sheet_workbook(
    headers: List[str], rows: List[List], sheet_title: str, output: io.BytesIO
) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = _sanitize_sheet_title(sheet_title, set())
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(output)


def _collect_sponsor_lookup(
    zip_payloads: List[bytes], legislator_name: str
) -> Tuple[dict[Tuple[str, str], str], dict[Tuple[str, str], dict], str]:
    sponsor_lookup: dict[Tuple[str, str], str] = {}
    sponsor_metadata: dict[Tuple[str, str], dict] = {}
    legislator_party_label: str = ""
    base_dirs = _get_payload_base_dirs(zip_payloads)
    try:
        csv_dirs = gather_session_csv_dirs(base_dirs)
    except FileNotFoundError:
        return sponsor_lookup, sponsor_metadata, legislator_party_label
    for csv_dir in csv_dirs:
        people_path = csv_dir / "people.csv"
        sponsors_path = csv_dir / "sponsors.csv"
        bills_path = csv_dir / "bills.csv"
        rollcalls_path = csv_dir / "rollcalls.csv"
        if not people_path.exists() or not sponsors_path.exists():
            continue

        target_id: Optional[int] = None
        target_party_code = ""
        with people_path.open(encoding="utf-8") as fh:
            reader = csv.DictReader(fh)
            for row in reader:
                name = (row.get("name") or "").strip()
                if name == legislator_name:
                    try:
                        target_id = int(row.get("people_id", ""))
                    except (TypeError, ValueError):
                        target_id = None
                    target_party_code = row.get("party", "")
                    break
        if target_id is None:
            continue
        if not legislator_party_label:
            legislator_party_label = _normalize_party_label(target_party_code)

        bill_map: dict[int, dict] = {}
        if bills_path.exists():
            with bills_path.open(encoding="utf-8") as fh:
                reader = csv.DictReader(fh)
                for row in reader:
                    try:
                        bill_id = int(row.get("bill_id", ""))
                    except (TypeError, ValueError):
                        continue
                    bill_map[bill_id] = row

        rollcalls_by_bill: dict[int, List[dict]] = {}
        if rollcalls_path.exists():
            with rollcalls_path.open(encoding="utf-8") as fh:
                reader = csv.DictReader(fh)
                for row in reader:
                    try:
                        bill_id = int(row.get("bill_id", ""))
                    except (TypeError, ValueError):
                        continue
                    rollcalls_by_bill.setdefault(bill_id, []).append(row)

        with sponsors_path.open(encoding="utf-8") as fh:
            reader = csv.DictReader(fh)
            for row in reader:
                try:
                    bill_id = int(row.get("bill_id", ""))
                    people_id = int(row.get("people_id", ""))
                except (TypeError, ValueError):
                    continue
                if people_id != target_id:
                    continue
                bill_row = bill_map.get(bill_id)
                if not bill_row:
                    continue
                session_id = str(bill_row.get("session_id", "")).strip()
                bill_number = str(bill_row.get("bill_number", "")).strip()
                if not bill_number:
                    continue
                key = (session_id, bill_number)
                position_value = row.get("position", "")
                try:
                    position_int = int(position_value)
                except (TypeError, ValueError):
                    position_int = None
                status = (
                    "Primary Sponsor"
                    if position_int == 1
                    else "Cosponsor"
                )
                existing = sponsor_lookup.get(key)
                if existing == "Primary Sponsor" and status != "Primary Sponsor":
                    continue
                if status == "Primary Sponsor" or key not in sponsor_lookup:
                    sponsor_lookup[key] = status
                roll_list = rollcalls_by_bill.get(bill_id, [])
                if key not in sponsor_metadata:
                    sponsor_metadata[key] = _build_sponsor_metadata(
                        bill_row, roll_list, status
                    )
                else:
                    if (
                        sponsor_metadata[key].get("sponsorship_status") != "Primary Sponsor"
                        or status == "Primary Sponsor"
                    ):
                        sponsor_metadata[key]["sponsorship_status"] = status
                    if roll_list and not sponsor_metadata[key].get("roll_call_id"):
                        sponsor_metadata[key].update(
                            _build_sponsor_metadata(bill_row, roll_list, status)
                        )
    return sponsor_lookup, sponsor_metadata, legislator_party_label


def _validate_archive_payload(payload: bytes) -> None:
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        with zipfile.ZipFile(io.BytesIO(payload)) as zf:
            zf.extractall(tmp_path)
        csv_dirs = gather_session_csv_dirs([tmp_path])
        if not csv_dirs:
            raise ValueError("Archive does not contain any LegiScan session directories.")
        for csv_dir in csv_dirs:
            missing = [name for name in REQUIRED_ARCHIVE_FILES if not (csv_dir / name).exists()]
            if missing:
                missing_list = ", ".join(missing)
                raise ValueError(f"Archive is missing required files: {missing_list} in {csv_dir.name}.")


def _save_uploaded_archive(filename: str, payload: bytes) -> Tuple[str, bool]:
    if not ARCHIVE_NAME_PATTERN.match(filename):
        raise ValueError("Filename must match pattern 'XX_Description.zip'.")
    _validate_archive_payload(payload)
    LOCAL_ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
    target_path = LOCAL_ARCHIVE_DIR / filename
    replaced = False
    if target_path.exists():
        new_status = _archive_session_status(filename)
        existing_status = _archive_session_status(target_path.name)
        if new_status != "current" or existing_status != "current":
            raise FileExistsError("Archive already exists in bulkLegiData.")
        try:
            existing_payload = target_path.read_bytes()
        except OSError as exc:
            raise FileExistsError(f"Unable to read existing archive: {exc}") from exc
        new_freshness = _infer_payload_freshness(filename, payload)
        existing_freshness = _infer_payload_freshness(target_path.name, existing_payload)
        if new_freshness and existing_freshness:
            if new_freshness <= existing_freshness:
                raise ArchiveUpToDate(filename)
            replaced = True
        elif new_freshness and not existing_freshness:
            replaced = True
        else:
            raise ArchiveUpToDate(filename)
    target_path.write_bytes(payload)
    return filename, replaced


def _upload_archives_to_github(saved_archive_names: List[str]) -> Tuple[bool, str]:
    if not saved_archive_names:
        return True, "No archives to upload."

    gh_cfg = st.secrets.get("github")
    if not gh_cfg:
        return False, "GitHub configuration missing in secrets."

    required_keys = ("token", "owner", "repo")
    missing_keys = [key for key in required_keys if not gh_cfg.get(key)]
    if missing_keys:
        return False, f"GitHub configuration missing keys: {', '.join(missing_keys)}"

    token = gh_cfg["token"]
    owner = gh_cfg["owner"]
    repo = gh_cfg["repo"]
    branch = gh_cfg.get("branch", "main")
    target_dir = gh_cfg.get("target_dir", "").strip().strip("/")

    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    }

    uploaded: List[str] = []
    errors: List[str] = []

    for filename in saved_archive_names:
        archive_path = LOCAL_ARCHIVE_DIR / filename
        if not archive_path.exists():
            errors.append(f"{filename}: file not found after saving.")
            continue

        try:
            file_bytes = archive_path.read_bytes()
        except OSError as exc:
            errors.append(f"{filename}: unable to read file ({exc}).")
            continue

        content_b64 = base64.b64encode(file_bytes).decode("utf-8")
        relative_path = f"{target_dir}/{filename}" if target_dir else filename
        api_url = f"https://api.github.com/repos/{owner}/{repo}/contents/{relative_path}"

        existing_sha = None
        probe = requests.get(api_url, headers=headers, params={"ref": branch}, timeout=20)
        if probe.status_code == 200:
            existing_sha = probe.json().get("sha")
        elif probe.status_code not in (200, 404):
            errors.append(
                f"{filename}: GitHub lookup failed ({probe.status_code}) -> {probe.text}"
            )
            continue

        payload = {
            "message": f"Add LegiScan archive {filename}",
            "content": content_b64,
            "branch": branch,
        }
        if existing_sha:
            payload["sha"] = existing_sha

        response = requests.put(api_url, headers=headers, json=payload, timeout=20)
        if response.status_code not in (200, 201):
            errors.append(
                f"{filename}: GitHub upload failed ({response.status_code}) -> {response.text}"
            )
            continue

        uploaded.append(filename)

    if errors:
        return False, "; ".join(errors)

    return True, f"Uploaded {len(uploaded)} archive(s) to GitHub."


def safe_int(value):
    try:
        return int(value)
    except (TypeError, ValueError):
        return 0


def prepare_summary_dataframe(rows: List[List]) -> pd.DataFrame:
    summary_df = pd.DataFrame(rows, columns=WORKBOOK_HEADERS)
    summary_df["Date_dt"] = pd.to_datetime(summary_df["Date"], errors="coerce")
    summary_df["Year"] = summary_df["Date_dt"].dt.year.astype("Int64")
    return summary_df


def _format_result_text(result_value: object) -> str:
    numeric = safe_int(result_value)
    if numeric == 1:
        return "Passed"
    if numeric == 0:
        return "Did not pass"
    return "Result unknown"


def _format_latest_action(meta: Dict[str, str]) -> str:
    last_action = (meta or {}).get("last_action") or (meta or {}).get("status_desc") or ""
    last_action_date = (meta or {}).get("last_action_date") or (meta or {}).get("status_date") or ""
    if not last_action:
        return "Latest action unavailable."
    if last_action_date:
        return f"{last_action} ({last_action_date})"
    return last_action


def _build_bullet_summary_doc(
    rows: pd.DataFrame,
    legislator_name: str,
    filter_label: str,
    bill_metadata: Dict[Tuple[str, str], Dict[str, str]],
    state_label: str,
) -> io.BytesIO:
    working_rows = rows.copy()
    if "Date_dt" not in working_rows.columns or working_rows["Date_dt"].isna().any():
        working_rows["Date_dt"] = pd.to_datetime(
            working_rows.get("Date"), errors="coerce"
        )

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    heading_style = doc.styles["Heading 1"]
    heading_style.font.name = "Arial"
    heading_style.font.size = Pt(10)
    heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    doc.add_heading(f"{legislator_name} - {filter_label}", level=1)

    state_display = (state_label or '').strip() or 'State'
    if state_display == ALL_STATES_LABEL:
        state_display = 'State'

    if working_rows.empty:
        doc.add_paragraph('No records available for this selection.')
    else:
        for _, row in working_rows.iterrows():
            session_id = str(row.get('Session') or '').strip()
            bill_number = str(row.get('Bill Number') or '').strip()
            meta = bill_metadata.get((session_id, bill_number), {})

            vote_dt = row.get('Date_dt')
            if pd.isna(vote_dt):
                first_sentence_prefix = 'Date Unknown'
                second_sentence_prefix = 'On an unknown date'
                fallback_vote_date = 'Date unknown'
            else:
                first_sentence_prefix = f"{vote_dt.strftime('%B')} {vote_dt.strftime('%Y')}"
                second_sentence_prefix = f"In {vote_dt.strftime('%B %Y')}"
                fallback_vote_date = f"{vote_dt.month}/{vote_dt.day}/{str(vote_dt.year)[-2:]}"

            bill_motion = (row.get('Bill Motion') or '').strip()
            bill_description = (row.get('Bill Description') or '').strip()
            bill_title = (row.get('Bill Title') or meta.get('title') or '').strip()

            primary_reference = bill_number or bill_motion or 'the bill'
            if bill_motion and not bill_number:
                primary_reference = bill_motion

            status_code = str(
                (meta.get('status_code') if meta else None)
                or row.get('Status')
                or ''
            ).strip()
            last_action = (meta.get('last_action') or meta.get('status_desc') or '').strip()
            last_action_date = (
                (meta.get('last_action_date') or meta.get('status_date') or '').strip()
                or fallback_vote_date
                or 'Date unknown'
            )
            last_action_date = _format_docx_date(last_action_date)
            chamber = (row.get('Chamber') or meta.get('chamber') or '').strip() or 'Chamber'
            vote_bucket = row.get('Vote Bucket')
            vote_upper, vote_lower = _resolve_vote_phrases(vote_bucket)
            vote_lead = vote_upper.title()

            description_text = bill_description or meta.get('title') or ''
            description_clean = (description_text or '').strip().rstrip('.!?').strip()
            if not description_clean:
                description_clean = 'No description provided'

            title_fragment = ""
            compare_description = bill_description
            if bill_title and bill_title.lower() != compare_description.lower():
                title_fragment = f", {bill_title}"

            status_sentence = _compose_status_sentence(
                status_code,
                primary_reference,
                chamber,
                last_action,
                (meta or {}).get("latest_votes"),
            )

            vote_url = (row.get('URL') or '').strip()

            paragraph = doc.add_paragraph()
            paragraph.style = normal_style
            paragraph.paragraph_format.space_after = Pt(12)
            paragraph.paragraph_format.space_before = Pt(0)
            sponsorship_status = (row.get("Sponsorship Status") or row.get("Sponsorship") or "").strip()
            if filter_label == "Sponsored/Cosponsored Bills":
                normalized_status = sponsorship_status.lower()
                sponsor_upper = (
                    "SPONSORED"
                    if "primary" in normalized_status or not normalized_status
                    else "COSPONSORED"
                )
                sponsor_lower = sponsor_upper.lower()
                sponsor_reference = bill_number or primary_reference
                first_sentence = (
                    f"{first_sentence_prefix}: {legislator_name} "
                    f"{sponsor_upper.title()} {sponsor_reference}{title_fragment}."
                )
                second_sentence = (
                    f"{second_sentence_prefix}, {legislator_name} "
                    f"{sponsor_lower} {sponsor_reference}, "
                    f"\"{description_clean}.\""
                )
            else:
                first_sentence = (
                    f"{first_sentence_prefix}: {legislator_name} {vote_lead} "
                    f"{primary_reference}{title_fragment}."
                )
                second_sentence = (
                    f"{second_sentence_prefix}, {legislator_name} {vote_lower} {primary_reference}: "
                    f"\"{description_clean}.\""
                )

            bold_run = paragraph.add_run(first_sentence + ' ')
            bold_run.bold = True

            paragraph.add_run(second_sentence + ' ')
            paragraph.add_run(status_sentence + ' ')

            paragraph.add_run('[')
            paragraph.add_run(f"{state_display} {chamber}, ")
            paragraph.add_run(f"{bill_number or 'Unknown bill'}, ")
            date_text_raw = last_action_date or 'Date unknown'
            if date_text_raw.lower() != 'date unknown':
                date_text = _format_docx_date(date_text_raw)
            else:
                date_text = date_text_raw
            if vote_url and date_text.lower() != 'date unknown':
                _add_hyperlink(paragraph, vote_url, date_text)
            else:
                paragraph.add_run(date_text)
            paragraph.add_run(']')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def apply_filters(
    summary_df: pd.DataFrame,
    *,
    filter_mode: str,
    search_term: str = "",
    year_selection: Optional[List[int]] = None,
    party_focus_option: str = "Legislator's Party",
    minority_percent: int = 20,
    min_group_votes: int = 0,
    max_vote_diff: int = 5,
    comparison_person: Optional[str] = None,
    selected_legislator: Optional[str] = None,
    zip_payloads: Optional[List[bytes]] = None,
    sponsor_metadata: Optional[dict] = None,
    legislator_party_label: str = "",
) -> Tuple[pd.DataFrame, int]:
    df = summary_df.copy()

    if year_selection:
        df = df[df["Year"].isin(year_selection)].copy()

    if search_term:
        description_mask = df["Bill Description"].astype(str).str.contains(
            search_term, case=False, na=False
        )
        df = df[description_mask].copy()

    if df.empty:
        if search_term:
            raise ValueError(f"No vote records found matching '{search_term}'.")
        raise ValueError("No vote records found for the selected criteria.")

    if filter_mode == "Search By Term" and not search_term:
        raise ValueError("Enter a search term to use the 'Search By Term' vote type.")

    if filter_mode == "Skipped Votes":
        vote_text = df["Vote"].astype(str).str.strip().str.lower()
        skip_mask = ~(
            vote_text.str.startswith("yea")
            | vote_text.str.startswith("nay")
            | vote_text.str.startswith("aye")
        )
        df = df[skip_mask].copy()
        if df.empty:
            raise ValueError("No skipped votes found for the selected criteria.")

    if filter_mode == "Sponsored/Cosponsored Bills":
        if "Sponsorship Status" in df.columns:
            sponsor_mask_series = df["Sponsorship Status"].astype(str).str.strip()
        else:
            sponsor_mask_series = pd.Series([""] * len(df), index=df.index)
        sponsor_mask = sponsor_mask_series != ""
        df = df[sponsor_mask].copy()
        existing_keys: Set[Tuple[str, str]] = {
            (str(session).strip(), str(bill_number).strip())
            for session, bill_number in zip(df.get("Session", []), df.get("Bill Number", []))
        }
        extra_rows: List[dict] = []
        if sponsor_metadata and selected_legislator:
            extra_rows = _create_sponsor_only_rows(
                sponsor_metadata,
                existing_keys,
                selected_legislator,
                legislator_party_label,
            )
        if extra_rows:
            df = pd.concat([df, pd.DataFrame(extra_rows)], ignore_index=True)
        df["Sponsorship Status"] = df["Sponsorship Status"].fillna("").astype(str)
        if df.empty:
            raise ValueError(
                "No sponsored or co-sponsored bills found for the selected legislator."
            )

    df["Roll Call ID"] = pd.to_numeric(df["Roll Call ID"], errors="coerce").astype(
        "Int64"
    )

    if filter_mode in {"Votes With Person", "Votes Against Person"}:
        if not comparison_person:
            raise ValueError("Select a comparison legislator in the sidebar.")
        if comparison_person == selected_legislator:
            raise ValueError("Choose a different legislator for comparison.")
        if zip_payloads is None:
            raise ValueError("Comparison vote data is unavailable.")
        comparison_votes = _collect_person_votes_from_zips(
            zip_payloads, comparison_person
        )
        if not comparison_votes:
            raise ValueError(f"No vote records found for {comparison_person}.")

        def lookup_comparison(rcid):
            if pd.isna(rcid):
                return pd.Series({"Comparison Vote": "", "Comparison Vote Bucket": ""})
            info = comparison_votes.get(int(rcid))
            if not info:
                return pd.Series({"Comparison Vote": "", "Comparison Vote Bucket": ""})
            return pd.Series(
                {
                    "Comparison Vote": info.get("vote_desc", ""),
                    "Comparison Vote Bucket": info.get("vote_bucket", ""),
                }
            )

        comparison_df = df["Roll Call ID"].apply(lookup_comparison)
        df = pd.concat([df, comparison_df], axis=1)
        df = df[df["Comparison Vote Bucket"] != ""].copy()
        if df.empty:
            verb = "with" if filter_mode == "Votes With Person" else "against"
            raise ValueError(
                f"No votes found where {selected_legislator} voted {verb} {comparison_person}."
            )

        main_bucket = df["Vote Bucket"]
        comp_bucket = df["Comparison Vote Bucket"]

        if filter_mode == "Votes With Person":
            comparison_mask = main_bucket == comp_bucket
        else:
            comparison_mask = (
                (main_bucket == "For") & (comp_bucket == "Against")
            ) | ((main_bucket == "Against") & (comp_bucket == "For"))

        df = df[comparison_mask].copy()
        if df.empty:
            verb = "with" if filter_mode == "Votes With Person" else "against"
            raise ValueError(
                f"No votes found where {selected_legislator} voted {verb} {comparison_person}."
            )

    def calc_metrics(row: pd.Series):
        bucket = row["Vote Bucket"]
        party = row["Person Party"]
        metrics = {
            "party_bucket_votes": None,
            "party_total_votes": None,
            "party_share": None,
            "chamber_bucket_votes": None,
            "chamber_total_votes": None,
            "chamber_share": None,
        }

        if party:
            party_bucket_col = f"{party}_{bucket}"
            party_total_col = f"{party}_Total"
            party_bucket = safe_int(row.get(party_bucket_col))
            party_total = safe_int(row.get(party_total_col))
            metrics["party_bucket_votes"] = party_bucket
            metrics["party_total_votes"] = party_total
            metrics["party_share"] = (
                party_bucket / party_total if party_total else None
            )

        chamber_bucket = safe_int(row.get(f"Total_{bucket}"))
        chamber_total = safe_int(row.get("Total_Total"))
        metrics["chamber_bucket_votes"] = chamber_bucket
        metrics["chamber_total_votes"] = chamber_total
        metrics["chamber_share"] = (
            chamber_bucket / chamber_total if chamber_total else None
        )
        return pd.Series(metrics)

    metrics_df = df.apply(calc_metrics, axis=1)
    df = pd.concat([df, metrics_df], axis=1)

    df["Person Party Display"] = df["Person Party"].map(PARTY_DISPLAY_MAP).fillna(
        df["Person Party"]
    )
    df["focus_party_label"] = df["Person Party Display"]
    df["focus_party_bucket_votes"] = df["party_bucket_votes"]
    df["focus_party_total_votes"] = df["party_total_votes"]
    df["focus_party_share"] = df["party_share"]

    focus_party_key = FOCUS_PARTY_LOOKUP.get(party_focus_option)
    if filter_mode == "Votes Against Party" and focus_party_key:
        focus_display_label = (
            "Independent" if focus_party_key == "Other" else party_focus_option
        )

        def calc_focus_metrics(row: pd.Series):
            bucket = row["Vote Bucket"]
            bucket_votes = safe_int(row.get(f"{focus_party_key}_{bucket}"))
            total_votes = safe_int(row.get(f"{focus_party_key}_Total"))
            share = bucket_votes / total_votes if total_votes else None
            return pd.Series(
                {
                    "focus_party_label": focus_display_label,
                    "focus_party_bucket_votes": bucket_votes,
                    "focus_party_total_votes": total_votes,
                    "focus_party_share": share,
                }
            )

        focus_metrics = df.apply(calc_focus_metrics, axis=1)
        df[
            [
                "focus_party_label",
                "focus_party_bucket_votes",
                "focus_party_total_votes",
                "focus_party_share",
            ]
        ] = focus_metrics

    deciding_condition = None
    if filter_mode == "Deciding Votes":
        total_for = df["Total_For"].apply(safe_int)
        total_against = df["Total_Against"].apply(safe_int)
        vote_diff = (total_for - total_against).abs()
        winning_bucket = pd.Series("Tie", index=df.index, dtype="object")
        winning_bucket = winning_bucket.mask(total_for > total_against, "For")
        winning_bucket = winning_bucket.mask(total_against > total_for, "Against")
        df["Vote Difference"] = vote_diff
        df["Winning Bucket"] = winning_bucket
        deciding_condition = (
            (vote_diff <= max_vote_diff)
            & winning_bucket.isin(["For", "Against"])
            & (df["Vote Bucket"] == winning_bucket)
        )

    apply_party_filter = filter_mode in {"Votes Against Party", "Minority Votes"}
    apply_chamber_filter = filter_mode == "Minority Votes"
    threshold_ratio = (
        minority_percent / 100.0 if (apply_party_filter or apply_chamber_filter) else None
    )
    min_votes = min_group_votes if (apply_party_filter or apply_chamber_filter) else 0

    filters = []
    if apply_party_filter:
        party_condition = (
            df["focus_party_share"].notna()
            & (df["focus_party_total_votes"] >= min_votes)
            & (df["focus_party_share"] <= threshold_ratio)
        )
        filters.append(party_condition)
    if apply_chamber_filter:
        chamber_condition = (
            df["chamber_share"].notna()
            & (df["chamber_total_votes"] >= min_votes)
            & (df["chamber_share"] <= threshold_ratio)
        )
        filters.append(chamber_condition)
    if filter_mode == "Deciding Votes" and deciding_condition is not None:
        filters.append(deciding_condition)

    pre_filter_count = len(df)

    if filters:
        mask = filters[0]
        for condition in filters[1:]:
            mask &= condition
        filtered_df = df[mask].copy()
    else:
        filtered_df = df.copy()

    dedupe_keys = [col for col in ["Roll Call ID", "Person"] if col in filtered_df.columns]
    if dedupe_keys:
        filtered_df = filtered_df.drop_duplicates(subset=dedupe_keys).reset_index(
            drop=True
        )

    if filtered_df.empty:
        if filter_mode == "Skipped Votes":
            raise ValueError("No skipped votes found for the selected criteria.")
        if filter_mode == "Votes Against Party":
            raise ValueError(
                "No votes found where the legislator sided with the specified minority."
            )
        if filter_mode == "Minority Votes":
            raise ValueError(
                "No votes found where the legislator and chamber were both in the minority."
            )
        if filter_mode == "Deciding Votes":
            raise ValueError(
                "No votes found within the specified deciding vote margin."
            )
        raise ValueError("No vote records found for the selected criteria.")

    return filtered_df, pre_filter_count


def write_multi_sheet_workbook(
    sheet_specs: List[Tuple[str, List[str], List[List]]], output: io.BytesIO
) -> None:
    wb = Workbook()
    first_sheet = True
    used_titles: Set[str] = set()
    for sheet_name, headers, rows in sheet_specs:
        safe_title = _sanitize_sheet_title(sheet_name, used_titles)
        if first_sheet:
            ws = wb.active
            ws.title = safe_title
            first_sheet = False
        else:
            ws = wb.create_sheet(title=safe_title)
        ws.append(headers)
        for row in rows:
            ws.append(row)
    wb.save(output)


def build_summary_dataframe(
    zip_payloads: List[bytes], legislator_name: Union[str, Iterable[str]]
) -> Tuple[pd.DataFrame, List[str]]:
    try:
        rows, missing_legislators = _collect_rows_from_zips(zip_payloads, legislator_name)
    except zipfile.BadZipFile:
        raise ValueError("One of the uploads could not be read as a ZIP archive.")
    except ValueError as exc:
        raise ValueError(str(exc)) from exc
    return prepare_summary_dataframe(rows), missing_legislators


def _collect_latest_action_date(zip_payloads: List[bytes]) -> Optional[dt.date]:
    latest_date: Optional[dt.date] = None
    base_dirs = _get_payload_base_dirs(zip_payloads)
    try:
        csv_dirs = gather_session_csv_dirs(base_dirs)
    except FileNotFoundError:
        return None
    for csv_dir in csv_dirs:
        bills_path = csv_dir / "bills.csv"
        if not bills_path.exists():
            continue
        with bills_path.open(encoding="utf-8") as fh:
            reader = csv.DictReader(fh)
            for row in reader:
                date_str = (row.get("last_action_date") or "").strip()
                if not date_str:
                    continue
                try:
                    parsed_date = dt.datetime.strptime(
                        date_str, "%Y-%m-%d"
                    ).date()
                except ValueError:
                    continue
                if latest_date is None or parsed_date > latest_date:
                    latest_date = parsed_date
    return latest_date


def _render_state_filter():
    st.sidebar.header("Filtering Parameters")
    state_label = st.sidebar.selectbox(
        "State",
        options=[ALL_STATES_LABEL] + [name for name, _ in STATE_CHOICES],
        index=0,
        key="state_filter_select",
        help="Filter archives by the state's two-letter prefix (e.g., MN_...).",
    )
    return state_label, STATE_NAME_TO_CODE.get(state_label)


st.set_page_config(page_title="LegiScan Vote Explorer", layout="wide")
st.title("LegiScan Vote Explorer")
st.caption(
    "Upload one or more LegiScan ZIP archives from the same state, then choose a legislator to generate a consolidated vote summary."
)

state_label, state_code = _render_state_filter()

uploaded_zips = st.file_uploader(
    "LegiScan ZIP file(s)", type="zip", accept_multiple_files=True
)

all_local_archive_paths = _list_local_archives()
local_archive_paths = [
    path
    for path in all_local_archive_paths
    if _archive_matches_state(path.name, state_code)
] if state_code else all_local_archive_paths
selected_local_archives: List[Path] = []
if local_archive_paths:
    local_lookup = {path.name: path for path in local_archive_paths}
    available_names = list(local_lookup.keys())
    existing_selection = st.session_state.get(BUNDLED_ARCHIVE_SESSION_KEY, [])
    filtered_selection = [name for name in existing_selection if name in local_lookup]
    if filtered_selection != existing_selection:
        st.session_state[BUNDLED_ARCHIVE_SESSION_KEY] = filtered_selection
    elif BUNDLED_ARCHIVE_SESSION_KEY not in st.session_state:
        st.session_state[BUNDLED_ARCHIVE_SESSION_KEY] = []

    select_all_col, clear_col = st.columns([1, 1])
    with select_all_col:
        if st.button("Select all bundled", width="stretch"):
            st.session_state[BUNDLED_ARCHIVE_SESSION_KEY] = available_names
    with clear_col:
        if st.button("Clear bundled", width="stretch"):
            st.session_state[BUNDLED_ARCHIVE_SESSION_KEY] = []

    selected_local_names = st.multiselect(
        "Bundled LegiScan archive(s)",
        options=available_names,
        key=BUNDLED_ARCHIVE_SESSION_KEY,
        format_func=_format_archive_option,
        help="Include ZIP archives stored in the repository (bulkLegiData).",
    )
    selected_local_archives = [
        local_lookup[name] for name in selected_local_names
    ]
else:
    if state_code and all_local_archive_paths:
        st.caption(
            f"No bundled archives match the selected state ({state_label})."
        )
    elif not uploaded_zips:
        st.caption(
            "Add additional archives under the 'bulkLegiData' directory to make them selectable here."
        )

if not uploaded_zips and not selected_local_archives:
    st.info("Upload ZIP files or select bundled archives to get started.")
    st.stop()

zip_payloads: List[bytes] = []
skipped_uploads: List[str] = []
duplicate_archives: List[str] = []
saved_archives: List[str] = []
updated_archives: List[str] = []
up_to_date_archives: List[str] = []
archive_save_errors: List[str] = []
seen_archive_keys: Set[str] = set()
processed_archive_names: List[str] = []
for uploaded_zip in uploaded_zips or []:
    if state_code and not _archive_matches_state(uploaded_zip.name, state_code):
        skipped_uploads.append(uploaded_zip.name)
        continue
    archive_key = _archive_key(uploaded_zip.name)
    if archive_key in seen_archive_keys:
        duplicate_archives.append(uploaded_zip.name)
        continue
    seen_archive_keys.add(archive_key)
    try:
        payload_bytes = uploaded_zip.getvalue()
        zip_payloads.append(payload_bytes)
        processed_archive_names.append(uploaded_zip.name)
    except Exception as exc:  # pragma: no cover - streamlit runtime guard
        st.error(f"Failed to read uploaded file '{uploaded_zip.name}': {exc}")
        st.stop()
    try:
        save_result = _save_uploaded_archive(uploaded_zip.name, payload_bytes)
        if save_result:
            saved_name, was_replaced = save_result
            if was_replaced:
                updated_archives.append(saved_name)
            else:
                saved_archives.append(saved_name)
    except ValueError as exc:
        archive_save_errors.append(f"{uploaded_zip.name}: {exc}")
    except zipfile.BadZipFile:
        archive_save_errors.append(f"{uploaded_zip.name}: Invalid ZIP archive.")
    except FileExistsError as exc:
        archive_save_errors.append(f"{uploaded_zip.name}: {exc}")
    except ArchiveUpToDate as exc:
        up_to_date_archives.append(exc.filename)
    except Exception as exc:  # pragma: no cover - unexpected
        archive_save_errors.append(f"{uploaded_zip.name}: {exc}")

for archive_path in selected_local_archives:
    archive_key = _archive_key(archive_path.name)
    if archive_key in seen_archive_keys:
        duplicate_archives.append(archive_path.name)
        continue
    seen_archive_keys.add(archive_key)
    try:
        zip_payloads.append(archive_path.read_bytes())
        processed_archive_names.append(archive_path.name)
    except OSError as exc:
        st.error(f"Failed to read bundled archive '{archive_path.name}': {exc}")
        st.stop()

archives_snapshot = tuple(sorted(processed_archive_names))

if processed_archive_names:
    status_groups = _group_archives_by_status(processed_archive_names)
    with st.sidebar.expander("Session status (from filenames)", expanded=False):
        for status_key, heading in [
            ("current", "Current"),
            ("upcoming", "Upcoming"),
            ("past", "Past"),
            ("unknown", "Unknown pattern"),
        ]:
            names = status_groups.get(status_key, [])
            if not names:
                continue
            display_limit = 5
            preview = ", ".join(names[:display_limit])
            if len(names) > display_limit:
                preview += f", +{len(names) - display_limit} more"
            st.markdown(f"**{heading} ({len(names)}):** {preview}")

if skipped_uploads:
    st.warning(
        f"Skipped uploads that do not match the selected state ({state_label}): "
        + ", ".join(skipped_uploads)
    )

all_changed_archives = saved_archives + updated_archives

if saved_archives:
    st.success(
        "Added new LegiScan archive(s) to 'bulkLegiData': "
        + ", ".join(saved_archives)
    )

if updated_archives:
    st.success(
        "Replaced existing archive(s) with newer data: "
        + ", ".join(updated_archives)
    )

if up_to_date_archives:
    st.info(
        "Archives already up to date: " + ", ".join(up_to_date_archives)
    )

if all_changed_archives:
    ok, github_message = _upload_archives_to_github(all_changed_archives)
    if ok:
        st.caption(github_message)
    else:
        st.warning(f"GitHub upload failed: {github_message}")

if archive_save_errors:
    st.warning(
        "Some uploads were not saved: " + "; ".join(archive_save_errors)
    )

if duplicate_archives:
    st.warning(
        "Skipped duplicate archives: " + ", ".join(duplicate_archives)
    )

if not zip_payloads:
    st.info("Provide at least one ZIP archive to continue.")
    st.stop()

latest_action_date: Optional[dt.date] = None
if state_code:
    latest_action_date = _collect_latest_action_date(zip_payloads)
    if latest_action_date:
        st.sidebar.caption(
            f"Latest bill action ({state_label}): "
            f"{latest_action_date.strftime('%B %d, %Y')}"
        )
    else:
        st.sidebar.caption(
            f"No bill action dates found for {state_label} in the selected archives."
        )

try:
    dataset_state, legislator_options = _collect_legislators_from_zips(zip_payloads)
except zipfile.BadZipFile:
    st.error("One or more uploads are not valid ZIP archives.")
    st.stop()
except FileNotFoundError as exc:
    st.error(f"{exc}")
    st.stop()
except ValueError as exc:
    st.error(str(exc))
    st.stop()

if not legislator_options:
    st.warning("No legislators found in the uploaded dataset.")
    st.stop()

if dataset_state:
    st.caption(f"Detected state: {dataset_state}")

year_options = _collect_years_from_zips(zip_payloads)

comparison_person = None
comparison_label = ""
max_vote_diff = 5

with st.sidebar:

    
    # Primary legislator (used for bullets, filenames, etc.)
    selected_legislator = st.selectbox("Legislator", legislator_options)

    # Additional legislators for group analysis
    extra_legislators = st.multiselect(
        "Additional legislators (optional)",
        legislator_options,
        default=[],
        help=(
            "Include more legislators in the current view. "
            "The primary legislator selected above will still be used "
            "for bullet summaries and filenames."
        ),
    )

    party_focus_option = "Legislator's Party"

    filter_mode = st.selectbox(
        "Vote type",
        options=[
            "All Votes",
            "Votes Against Party",
            "Votes With Person",
            "Votes Against Person",
            "Minority Votes",
            "Deciding Votes",
            "Sponsored/Cosponsored Bills",
            "Skipped Votes",
            "Search By Term",
        ],
        index=0,
        help="Choose a predefined view of the legislator's voting record.",
    )

    # NEW: initialize party_focus_option (prevents NameError)
    party_focus_option = "Legislator's Party"

    # NEW: initialize minority_percent & min_group_votes too (safe defaults)
    minority_percent = 20
    min_group_votes = 0

    # Define search_term
    search_term = st.text_input("Search term (bill description)", value="")

    # Build the list of all selected legislators for group operations
    selected_legislators = [selected_legislator] + [
        name for name in extra_legislators if name != selected_legislator
    ]


    if filter_mode == "All Votes":
        minority_percent = 20
        min_group_votes = 0
    elif filter_mode == "Votes Against Party":
        party_focus_option = st.selectbox(
            "Party voting against",
            options=["Legislator's Party", "Democrat", "Republican", "Independent"],
            index=0,
            key="votes_against_party_focus",
            help="Choose which party's vote breakdown to compare against.",
        )
        minority_percent = st.slider(
            "Minority threshold (%)",
            min_value=0,
            max_value=100,
            value=20,
            key="votes_against_party_threshold",
            help="Keep votes where the selected party supported the legislator's position at or below this percentage.",
        )
        min_group_votes = st.slider(
            "Minimum votes in group",
            min_value=0,
            max_value=200,
            value=5,
            key="votes_against_party_min_votes",
            help="Ignore vote records where the compared party cast fewer total votes than this threshold.",
        )
        st.caption("Shows votes where the legislator sided with a minority of the chosen party.")
    elif filter_mode == "Votes With Person":
        comparison_label = "Person voting with"
        comparison_person = st.selectbox(
            comparison_label,
            options=legislator_options,
            index=0,
            key="votes_with_person_select",
            help="Select another legislator to find votes where they aligned.",
        )
        minority_percent = 20
        min_group_votes = 0
        st.caption("Shows votes where the legislator and selected colleague cast the same vote.")
    elif filter_mode == "Votes Against Person":
        comparison_label = "Person voting against"
        comparison_person = st.selectbox(
            comparison_label,
            options=legislator_options,
            index=0,
            key="votes_against_person_select",
            help="Select another legislator to find votes where their positions opposed each other.",
        )
        minority_percent = 20
        min_group_votes = 0
        st.caption("Shows votes where the legislator and selected colleague took opposing sides.")
    elif filter_mode == "Minority Votes":
        minority_percent = st.slider(
            "Minority threshold (%)",
            min_value=0,
            max_value=100,
            value=20,
            key="minority_votes_threshold",
            help="Keep votes where the legislator's party supported their position at or below this percentage.",
        )
        min_group_votes = st.slider(
            "Minimum votes in group",
            min_value=0,
            max_value=200,
            value=5,
            key="minority_votes_min_votes",
            help="Ignore vote records where the compared group cast fewer total votes than this threshold.",
        )
        st.caption("Shows votes where the legislator sided with a minority of both their party and the full chamber.")
    elif filter_mode == "Deciding Votes":
        minority_percent = 20
        min_group_votes = 0
        max_vote_diff = st.slider(
            "Maximum votes difference",
            min_value=1,
            max_value=50,
            value=5,
            key="deciding_votes_max_diff",
            help="Limit to votes where the margin between Yeas and Nays is within this amount.",
        )
        st.caption("Shows votes where the legislator's side prevailed by the specified margin or less.")
    elif filter_mode == "Sponsored/Cosponsored Bills":
        minority_percent = 20
        min_group_votes = 0
        max_vote_diff = 5
        st.caption("Shows votes on bills the legislator sponsored or co-sponsored.")
    elif filter_mode == "Search By Term":
        minority_percent = 20
        min_group_votes = 0
        st.caption("Shows votes where the bill description matches the search term.")
    else:  # Skipped Votes
        minority_percent = 20
        min_group_votes = 0
        st.caption("Shows votes where the legislator did not cast a Yea or Nay.")

    year_selection = st.multiselect(
        "Year",
        options=year_options,
        default=year_options,
        help="Restrict votes to selected calendar years.",
    )

year_selection_snapshot = (
    tuple(sorted(int(year) for year in year_selection))
    if year_selection
    else ()
)
inputs_snapshot = {
    "selected_legislator": selected_legislator,
    # NEW: capture the full group so cached views refresh correctly
    "selected_legislators": tuple(selected_legislators),
    "filter_mode": filter_mode,
    "search_term": search_term,
    "year_selection": year_selection_snapshot,
    "party_focus_option": party_focus_option,
    "minority_percent": minority_percent,
    "min_group_votes": min_group_votes,
    "max_vote_diff": max_vote_diff,
    "comparison_person": comparison_person,
    "archives": archives_snapshot,
    "state_label": state_label,
    "state_code": state_code,
}

if not selected_legislator:
    st.stop()

generate_summary_clicked = st.button("Generate current view summary")
generate_workbook_clicked = st.button("Generate all views workbook")

summary_df: Optional[pd.DataFrame] = None
filtered_df: Optional[pd.DataFrame] = None
total_count: int = 0
sponsor_metadata: dict[Tuple[str, str], dict] = {}
legislator_party_label: str = ""
missing_legislators: List[str] = []


bill_metadata: Dict[Tuple[str, str], Dict[str, str]] = {}
view_state = st.session_state.get(VIEW_STATE_SESSION_KEY)
if (
    view_state
    and not generate_summary_clicked
    and not generate_workbook_clicked
    and view_state.get("inputs") == inputs_snapshot
):
    stored_summary = view_state.get("summary_df")
    stored_filtered = view_state.get("filtered_df")
    if isinstance(stored_summary, pd.DataFrame):
        summary_df = stored_summary.copy()
    if isinstance(stored_filtered, pd.DataFrame):
        filtered_df = stored_filtered.copy()
    total_count = view_state.get("total_count", len(filtered_df) if filtered_df is not None else 0)
    dataset_state = view_state.get("dataset_state", dataset_state)
    state_code = view_state.get("state_code", state_code)
    sponsor_metadata = view_state.get("sponsor_metadata", sponsor_metadata)
    bill_metadata = view_state.get("bill_metadata", bill_metadata)
    legislator_party_label = view_state.get("legislator_party_label", legislator_party_label)
if generate_summary_clicked or generate_workbook_clicked:
    spinner_label = (
        "Processing LegiScan data..."
        if generate_summary_clicked
        else "Compiling workbook across vote types..."
    )
    with st.spinner(spinner_label):
        try:
            summary_df, missing_legislators = build_summary_dataframe(
                zip_payloads, selected_legislators
            )
        except ValueError as exc:
            st.warning(str(exc))
            st.stop()

    if missing_legislators:
        st.warning(
            "No vote records found for: " + ", ".join(sorted(set(missing_legislators)))
        )

    sponsor_lookup, sponsor_metadata, legislator_party_label = _collect_sponsor_lookup(
        zip_payloads, selected_legislator
    )
    bill_metadata = _collect_bill_metadata(zip_payloads)
    
    # NEW: sponsor key sets for additional legislators
    extra_sponsor_sets: Dict[str, Set[Tuple[str, str]]] = {}
    for extra_name in extra_legislators:
        try:
            _, extra_metadata, _ = _collect_sponsor_lookup(zip_payloads, extra_name)
        except ValueError:
            # If a legislator has no sponsor metadata, treat as empty set
            extra_metadata = {}
        extra_sponsor_sets[extra_name] = set(extra_metadata.keys())
    
    session_series = (
        summary_df["Session"].astype(str)
        if "Session" in summary_df.columns
        else pd.Series([""] * len(summary_df))
    )
    bill_number_series = (
        summary_df["Bill Number"].astype(str)
        if "Bill Number" in summary_df.columns
        else pd.Series([""] * len(summary_df))
    )
    summary_df["Sponsorship Status"] = [
        sponsor_lookup.get((session, bill_number), "")
        for session, bill_number in zip(session_series, bill_number_series)
    ]

if generate_summary_clicked and summary_df is not None:
    try:
        filtered_df, total_count = apply_filters(
            summary_df,
            filter_mode=filter_mode,
            search_term=search_term,
            year_selection=year_selection,
            party_focus_option=party_focus_option,
            minority_percent=minority_percent,
            min_group_votes=min_group_votes,
            max_vote_diff=max_vote_diff,
            comparison_person=comparison_person,
            selected_legislator=selected_legislator,
            zip_payloads=zip_payloads,
            sponsor_metadata=sponsor_metadata,
            legislator_party_label=legislator_party_label,
        )
    except ValueError as exc:
        st.warning(str(exc))
        st.stop()
        
    # NEW: if a group of legislators is selected and we're in the
    # "Sponsored/Cosponsored Bills" view, keep only bills that are
    # sponsored/cosponsored by ALL selected legislators.
    if (
        filter_mode == "Sponsored/Cosponsored Bills"
        and extra_legislators
        and filtered_df is not None
        and not filtered_df.empty
    ):
        # Keys for primary legislator's sponsored/cosponsored bills
        primary_keys: Set[Tuple[str, str]] = set(sponsor_metadata.keys())

        # Collect sets for all legislators in the group
        group_sets: List[Set[Tuple[str, str]]] = []
        if primary_keys:
            group_sets.append(primary_keys)
        for extra_name in extra_legislators:
            group_sets.append(extra_sponsor_sets.get(extra_name, set()))

        # If any set is empty, intersection will be empty → early exit
        non_empty_sets = [s for s in group_sets if s]
        if not non_empty_sets:
            st.warning(
                "No bills found that are sponsored or co-sponsored by all "
                "of the selected legislators."
            )
            st.stop()

        shared_keys = set.intersection(*non_empty_sets)
        if not shared_keys:
            st.warning(
                "No bills found that are sponsored or co-sponsored by all "
                "of the selected legislators."
            )
            st.stop()

        # Attach a bill key to each row and filter to the shared set
        filtered_df = filtered_df.copy()
        filtered_df["_bill_key"] = list(
            zip(
                filtered_df["Session"].astype(str),
                filtered_df["Bill Number"].astype(str),
            )
        )
        filtered_df = filtered_df[filtered_df["_bill_key"].isin(shared_keys)].copy()
        filtered_df = filtered_df.drop(columns=["_bill_key"])
        total_count = len(filtered_df)

        if filtered_df.empty:
            st.warning(
                "No bills found that are sponsored or co-sponsored by all "
                "of the selected legislators."
            )
            st.stop()


    row_key_series = filtered_df.apply(_compose_row_key, axis=1).astype(str)
    filtered_df = filtered_df.assign(_row_key=row_key_series)
    st.session_state[VIEW_STATE_SESSION_KEY] = {
        "inputs": inputs_snapshot.copy(),
        "summary_df": summary_df.copy(),
        "filtered_df": filtered_df.copy(),
        "total_count": total_count,
        "dataset_state": dataset_state,
        "state_code": state_code,
        "sponsor_metadata": sponsor_metadata,
        "bill_metadata": bill_metadata,
        "legislator_party_label": legislator_party_label,
    }
elif filtered_df is not None and summary_df is not None:
    if "_row_key" not in filtered_df.columns:
        row_key_series = filtered_df.apply(_compose_row_key, axis=1).astype(str)
        filtered_df = filtered_df.assign(_row_key=row_key_series)
        cached_view = st.session_state.get(VIEW_STATE_SESSION_KEY)
        if cached_view and cached_view.get("inputs") == inputs_snapshot:
            cached_view["filtered_df"] = filtered_df.copy()
            st.session_state[VIEW_STATE_SESSION_KEY] = cached_view

if filtered_df is not None and summary_df is not None:
    bullet_state_key = f"bullet_flags::{selected_legislator}::{filter_mode}"
    row_key_series = filtered_df["_row_key"].astype(str)
    person_series = filtered_df.get("Person", pd.Series([""] * len(filtered_df))).astype(str)
    primary_mask = person_series == str(selected_legislator)
    primary_row_keys = set(row_key_series[primary_mask])
    stored_flags_raw = {
        str(key): bool(value)
        for key, value in st.session_state.get(bullet_state_key, {}).items()
        if str(key) in primary_row_keys
    }
    for key in primary_row_keys:
        stored_flags_raw.setdefault(key, False)
    st.session_state[bullet_state_key] = stored_flags_raw
    existing_bullet_flags = stored_flags_raw
    primary_filtered_df = filtered_df[primary_mask].copy()

    filtered_count = len(filtered_df)
    state_label = f" ({dataset_state})" if dataset_state else ""
    group_names = [selected_legislator] + extra_legislators
    group_label = ", ".join(dict.fromkeys(name for name in group_names if name))
    st.success(
        f"Compiled {total_count} votes for {group_label}{state_label}. "
        f"Showing {filtered_count} after filters."
    )

    export_source_df = filtered_df.copy()
    if filter_mode == "Sponsored/Cosponsored Bills":
        export_source_df = _prepare_sponsor_dataframe(export_source_df)
        export_headers = [
            header for header in WORKBOOK_HEADERS if header not in SPONSOR_DROP_COLUMNS
        ]
    else:
        export_headers = list(WORKBOOK_HEADERS)

    if "_row_key" in export_source_df.columns:
        export_source_df = export_source_df.drop(columns=["_row_key"])

    if "Person" in export_headers:
        person_index = export_headers.index("Person") + 1
    else:
        person_index = len(export_headers)
    if "Sponsorship" not in export_headers:
        export_headers.insert(person_index, "Sponsorship")

    export_df = (
        export_source_df.rename(columns={"Sponsorship Status": "Sponsorship"})
        .reindex(columns=export_headers)
        .fillna("")
    )
    export_rows = export_df.values.tolist()

    download_buffer = io.BytesIO()
    _write_single_sheet_workbook(export_headers, export_rows, selected_legislator, download_buffer)
    download_buffer.seek(0)

    download_filename = _make_download_filename(
        selected_legislator,
        filter_mode,
        dataset_state=dataset_state,
        fallback_state=state_code,
    )
    st.download_button(
        label="Download filtered Excel sheet",
        data=download_buffer.getvalue(),
        file_name=download_filename,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    primary_filtered_count = len(primary_filtered_df)

    if filtered_count > 0 and primary_filtered_count > 0:
        state_display_value = state_label
        if state_display_value == ALL_STATES_LABEL:
            state_display_value = dataset_state or state_code or ""
        allowed_keys = {key for key, value in existing_bullet_flags.items() if value}
        if allowed_keys:
            bullet_source_df = primary_filtered_df[
                primary_filtered_df["_row_key"].isin(allowed_keys)
            ].copy()
        else:
            bullet_source_df = primary_filtered_df.copy()
        if filter_mode == "Sponsored/Cosponsored Bills":
            bullet_rows = _prepare_sponsor_dataframe(bullet_source_df, drop_date=False)
        else:
            bullet_rows = bullet_source_df.copy()
        if "_row_key" in bullet_rows.columns:
            bullet_rows = bullet_rows.drop(columns=["_row_key"])
        bullet_doc_buffer = _build_bullet_summary_doc(
            bullet_rows,
            selected_legislator,
            filter_mode,
            bill_metadata,
            state_display_value,
        )
        bullet_filename = _make_docx_filename(
            selected_legislator,
            filter_mode,
            dataset_state=dataset_state,
            fallback_state=state_code,
        )
        st.download_button(
            label="Download bullet summary",
            data=bullet_doc_buffer.getvalue(),
            file_name=bullet_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_bullet_summary",
        )
    elif filtered_count == 0:
        st.info("No records available for bullet summary.")
    else:
        st.info(
            "The selected legislator has no records in this view, so the bullet summary download is disabled."
        )

    display_df = filtered_df.copy()
    if "Result" in display_df.columns:
        display_df["Result"] = (
            pd.to_numeric(display_df["Result"], errors="coerce").astype("Int64")
        )
    display_df["Date"] = display_df["Date_dt"].dt.date
    display_df["Legislator Party"] = display_df["Person Party"].map(
        PARTY_DISPLAY_MAP
    ).fillna(display_df["Person Party"])
    display_df["Focus Party"] = display_df["focus_party_label"]
    if filter_mode in {"Votes With Person", "Votes Against Person"}:
        display_df["Comparison Legislator"] = comparison_person
    display_df["Legislator Party Minority %"] = (
        display_df["party_share"] * 100
    ).round(1)
    display_df["Focus Party Minority %"] = (
        display_df["focus_party_share"] * 100
    ).round(1)
    display_df["Chamber Minority %"] = (
        display_df["chamber_share"] * 100
    ).round(1)
    display_df = display_df.rename(
        columns={
            "party_bucket_votes": "Legislator Party Votes (same position)",
            "party_total_votes": "Legislator Party Total Votes",
            "focus_party_bucket_votes": "Focus Party Votes (same position)",
            "focus_party_total_votes": "Focus Party Total Votes",
            "chamber_bucket_votes": "Chamber Votes (same position)",
            "chamber_total_votes": "Chamber Total Votes",
            "Vote Difference": "Vote Margin",
            "Winning Bucket": "Winning Side",
            "Sponsorship Status": "Sponsorship",
        }
    )
    if "Sponsorship" in display_df.columns and "Person" in display_df.columns:
        column_order = list(display_df.columns)
        sponsorship_index = column_order.index("Sponsorship")
        person_index = column_order.index("Person")
        if sponsorship_index != person_index + 1:
            sponsorship_column = column_order.pop(sponsorship_index)
            column_order.insert(person_index + 1, sponsorship_column)
            display_df = display_df[column_order]

    st.subheader("Vote Breakdown")
    display_table = display_df.drop(
        columns=[
            "party_share",
            "focus_party_share",
            "chamber_share",
            "focus_party_label",
            "Person Party Display",
            "Date_dt",
        ],
        errors="ignore",
    )
    if "_row_key" not in display_table.columns:
        display_table["_row_key"] = filtered_df["_row_key"]
    row_keys = display_table["_row_key"].astype(str)
    bullet_defaults = row_keys.map(existing_bullet_flags).fillna(False).astype(bool)
    bullet_defaults_list = bullet_defaults.tolist()
    display_table = display_table.drop(columns=["_row_key"])
    display_table.insert(
        0,
        "BULLET?",
        bullet_defaults_list,
    )
    display_table.index = row_keys
    form_key = f"bullet_toggle_form::{selected_legislator}::{filter_mode}"
    editor_key = f"vote_breakdown_editor::{selected_legislator}::{filter_mode}"
    column_configs: Dict[str, st.column_config.Column] = {
        "BULLET?": st.column_config.CheckboxColumn(
            label="BULLET?",
            help="Toggle to include this vote in the bullet summary download.",
            default=True,
        )
    }
    for col in display_table.columns:
        if col == "BULLET?":
            continue
        column_configs[col] = st.column_config.Column(disabled=True)
    with st.form(form_key):
        edited_table = st.data_editor(
            display_table,
            column_config=column_configs,
            hide_index=True,
            width="stretch",
            height=600,
            key=editor_key,
        )
        apply_bullet_changes = st.form_submit_button("Update bullet selections")
    if apply_bullet_changes:
        updated_flags = {}
        for idx, val in edited_table["BULLET?"].astype(bool).to_dict().items():
            if str(idx) in primary_row_keys:
                updated_flags[str(idx)] = bool(val)
        st.session_state[bullet_state_key] = updated_flags
        st.success("Bullet selections updated for the current view.")
    else:
        st.caption(
            "Toggle the checkboxes for the primary legislator's rows and click **Update bullet selections** to refresh the bullet summary download."
        )

if generate_workbook_clicked and summary_df is not None:
    stored_votes_against_focus = st.session_state.get(
        "votes_against_party_focus", "Legislator's Party"
    )
    stored_votes_against_threshold = st.session_state.get(
        "votes_against_party_threshold", 20
    )
    stored_votes_against_min_votes = st.session_state.get(
        "votes_against_party_min_votes", 5
    )
    stored_minority_threshold = st.session_state.get(
        "minority_votes_threshold", 20
    )
    stored_minority_min_votes = st.session_state.get(
        "minority_votes_min_votes", 5
    )
    stored_deciding_max_diff = st.session_state.get(
        "deciding_votes_max_diff", 5
    )

    workbook_views = [
        ("All Votes", {}),
        (
            "Votes Against Party",
            {
                "party_focus_option": stored_votes_against_focus,
                "minority_percent": stored_votes_against_threshold,
                "min_group_votes": stored_votes_against_min_votes,
            },
        ),
        (
            "Minority Votes",
            {
                "minority_percent": stored_minority_threshold,
                "min_group_votes": stored_minority_min_votes,
            },
        ),
        (
            "Deciding Votes",
            {
                "max_vote_diff": stored_deciding_max_diff,
            },
        ),
        ("Sponsored/Cosponsored Bills", {}),
        ("Skipped Votes", {}),
    ]
    sheet_rows: List[Tuple[str, List[str], List[List]]] = []
    empty_views: List[str] = []
    base_params = {
        "search_term": "",
        "year_selection": None,
        "party_focus_option": "Legislator's Party",
        "minority_percent": 20,
        "min_group_votes": 0,
        "max_vote_diff": 5,
        "comparison_person": None,
        "selected_legislator": selected_legislator,
        "zip_payloads": zip_payloads,
        "sponsor_metadata": sponsor_metadata,
        "legislator_party_label": legislator_party_label,
    }

    for sheet_name, overrides in workbook_views:
        params = base_params.copy()
        params.update(overrides)
        try:
            sheet_df, _ = apply_filters(
                summary_df, filter_mode=sheet_name, **params
            )
        except ValueError:
            empty_views.append(sheet_name)
            sheet_df = summary_df.iloc[0:0].copy()
        sheet_source_df = sheet_df
        sheet_headers = list(WORKBOOK_HEADERS)
        if sheet_name == "Sponsored/Cosponsored Bills":
            sheet_source_df = _prepare_sponsor_dataframe(sheet_df)
            sheet_headers = [
                header for header in sheet_headers if header not in SPONSOR_DROP_COLUMNS
            ]

        if "Person" in sheet_headers:
            person_idx = sheet_headers.index("Person") + 1
        else:
            person_idx = len(sheet_headers)
        sponsorship_header = "Sponsorship"
        if sponsorship_header not in sheet_headers:
            sheet_headers.insert(person_idx, sponsorship_header)
        sheet_df_export = (
            sheet_source_df.rename(columns={"Sponsorship Status": sponsorship_header})
            .reindex(columns=sheet_headers)
            .fillna("")
        )
        sheet_data = sheet_df_export.values.tolist()
        sheet_rows.append((sheet_name, sheet_headers, sheet_data))

    workbook_buffer = io.BytesIO()
    write_multi_sheet_workbook(sheet_rows, workbook_buffer)
    workbook_buffer.seek(0)

    st.success("Compiled vote summary workbook across key views.")
    workbook_filename = _make_download_filename(
        selected_legislator,
        "FULL",
        dataset_state=dataset_state,
        fallback_state=state_code,
    )

    st.download_button(
        label="Download vote summary workbook",
        data=workbook_buffer.getvalue(),
        file_name=workbook_filename,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="download_multi_view_workbook",
    )

    if empty_views:
        st.info(
            "No data available for: " + ", ".join(empty_views)
        )
